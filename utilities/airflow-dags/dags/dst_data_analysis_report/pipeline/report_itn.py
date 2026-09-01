"""
report_itn.py — ITN/LLIN performance Excel -> Word .docx + Slack text

Separate module from report.py by design (see analyze_itn.py's docstring for the
full reasoning): report.py's _build_doc() hardcodes "Children Treated" / drug1-drug2
splits and individual-child DQ metrics throughout nearly every section, not just in
a couple of branches — retrofitting ITN there would mean wrapping most of a 300-line
function in conditionals. Shares only the generic docx styling helpers in
dst_data_analysis_report.pipeline.core.word — report.py itself is never imported and its SPAQ/AZM path
is untouched.

Section structure deliberately mirrors report.py's Word report (same rigor, same
sections) so the two report types read consistently to the same audience:
  1. Campaign Overview
  2. Program Issues (Claude/Groq narrative, same as SPAQ/AZM)
  3. Distribution Data Analysis (by Province/LGA, by Facility)
  4. Data Quality Summary (ITN-appropriate DQ set — see analyze_itn.py's DQ mapping)
  5. Conclusion
"""
import json
import logging
import os
from collections import defaultdict
from datetime import datetime, timedelta

import openpyxl
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dst_data_analysis_report.pipeline.core.llm import generate_narrative
from dst_data_analysis_report.pipeline.core.word import (
    ALT_FILL, FONT, GREY_RGB, STATUS_COLOR, TITLE_RGB,
    add_heading, add_hyperlink, add_para, cov_band, dat, hdr, set_cell_bg,
    set_cell_borders, two_col_table,
)

log = logging.getLogger(__name__)


# ── data loader ────────────────────────────────────────────────────────────────

def _load_perf_itn(path):
    """
    Read analyze_itn.py's ALL LGAS tab (LGA-grain rows — the grain that actually
    carries a real target, see analyze_itn.py's docstring) and its FACILITY DETAIL
    tab (facility-grain, no target/coverage/status — none exists at that grain).
    """
    wb = openpyxl.load_workbook(path, read_only=True)
    if "ALL LGAS" not in wb.sheetnames:
        wb.close()
        raise FileNotFoundError(f"'ALL LGAS' tab missing in {path} — analyze_itn.py may have failed")
    ws = wb["ALL LGAS"]
    rows_raw = [
        r for r in ws.iter_rows(min_row=3, values_only=True)
        if r[2] and str(r[2]).strip() not in ("", "GRAND TOTAL")
    ]
    wb.close()

    # col order (see analyze_itn.HEADERS):
    # #, Province, LGA, Facilities, Target HH, HH Visited, HH Cov%,
    # Target Pop, Pop Covered, Pop Cov%, Target ITN, Nets Distributed, ITN Cov%,
    # Status, Records, Dup Records, Missing HH Head, Missing GPS,
    # Manual Codes, Scanned Codes, % Scanned, Missing Codes,
    # then (appended at the END — cols 22-25, may be ABSENT on pre-matrix files)
    # the duplicate-matrix buckets in _DUP_MATRIX_KEYS order.
    lga_d = {}

    for row in rows_raw:
        # Pad short rows (older-generation files) so the fixed unpack below
        # cannot raise — absent trailing columns read as None.
        if len(row) < 22:
            row = tuple(row) + (None,) * (22 - len(row))
        (_, province, lga, facs, hh_t, hh_v, hh_cov, pop_t, pop_c, pop_cov,
         net_t, net_d, net_cov, status, records, dup, miss_hh, miss_gps,
         manual_codes, scanned_codes, pct_scanned, miss_codes) = row[:22]

        def i(v): return int(v or 0)

        def dm(idx):
            # Duplicate-matrix cell: absent column (old file) or empty cell
            # (matrix off/failed that run) -> None ("not measured"), never a
            # fabricated 0 — the doc section is omitted entirely on None.
            v = row[idx] if len(row) > idx else None
            if v is None or str(v).strip() == "":
                return None
            try:
                return int(float(v))
            except (ValueError, TypeError):
                return None

        if not lga:
            continue
        lga_d[str(lga).strip()] = dict(
            province=str(province or "").strip(), facs=i(facs),
            hh_target=i(hh_t), hh_visited=i(hh_v),
            pop_target=i(pop_t), pop_covered=i(pop_c),
            net_target=i(net_t), nets_distributed=i(net_d),
            status=str(status or "").strip(),
            records=i(records), dup_records=i(dup),
            missing_hh_head=i(miss_hh), missing_gps=i(miss_gps),
            manual_codes=i(manual_codes), scanned_codes=i(scanned_codes),
            missing_codes=i(miss_codes),
            **{k: dm(22 + n) for n, k in enumerate(_DUP_MATRIX_KEYS)},
        )

    facilities = _load_facility_detail_itn(path)
    return lga_d, facilities


def _load_facility_detail_itn(path):
    """
    Read analyze_itn.py's FACILITY DETAIL tab — raw records/DQ counts per
    facility, no target/coverage/status columns (none exist at facility grain).
    Used only for the Low Activity Facilities list.
    """
    wb = openpyxl.load_workbook(path, read_only=True)
    if "FACILITY DETAIL" not in wb.sheetnames:
        wb.close()
        return []
    ws = wb["FACILITY DETAIL"]
    rows_raw = [
        r for r in ws.iter_rows(min_row=3, values_only=True)
        if r[3] and str(r[3]).strip() not in ("", "GRAND TOTAL")
    ]
    wb.close()

    # col order (see analyze_itn.FACILITY_HEADERS):
    # #, Province, LGA, Health Facility, Records, Dup Records, Households Visited,
    # Nets Distributed, Population Covered, Missing HH Head, Missing GPS,
    # Manual Codes, Scanned Codes, % Scanned, Missing Codes
    facilities = []
    for row in rows_raw:
        (_, province, lga, fac, records, dup, hh_v, net_d, pop_c,
         miss_hh, miss_gps, manual_codes, scanned_codes, pct_scanned, miss_codes) = row[:15]

        def i(v): return int(v or 0)

        facilities.append({
            "province": str(province or "").strip(), "lga": str(lga or "").strip(),
            "fac": str(fac or "").strip(), "records": i(records),
            "dup_records": i(dup), "households_visited": i(hh_v),
            "nets_distributed": i(net_d), "population_covered": i(pop_c),
        })
    return facilities


def _grand_totals(lga_d):
    g = defaultdict(int)
    for D in lga_d.values():
        for k, v in D.items():
            if isinstance(v, (int, float)):
                g[k] += v
    return g


# Same key order as analyze_itn._DUP_KEYS / its appended _DUP_HEADERS columns —
# kept in sync manually since this module reads the Excel, not analyze_itn.
_DUP_MATRIX_KEYS = ("dup_su_sd", "dup_su_dd", "dup_du_sd", "dup_du_dd")


def _dup_matrix_totals(lga_d):
    """
    Totals of the duplicate-distribution matrix, or None when the matrix was
    not measured for this extract (dup_matrix off, enrichment failed, or a
    pre-matrix Excel) — callers render nothing on None, never fabricated zeros.
    Measured/None is expected to be uniform across rows; a mix (hand-edited
    file, writer bug) is warned about since None rows would sum as zeros.
    """
    if not lga_d:
        return None
    none_rows = sum(1 for D in lga_d.values()
                    if all(D.get(k) is None for k in _DUP_MATRIX_KEYS))
    if none_rows == len(lga_d):
        return None
    if none_rows:
        log.warning(f"[report_itn] dup matrix: {none_rows}/{len(lga_d)} LGA rows unmeasured "
                    f"but others measured — totals treat unmeasured rows as 0")
    return {k: sum(D.get(k) or 0 for D in lga_d.values()) for k in _DUP_MATRIX_KEYS}


def _cov_str(numer, denom):
    if not denom:
        return "N/A"
    return f"{numer/denom*100:.1f}%"


def _load_sync_summary_itn(path):
    """
    Read cdd_sync_itn.py's SUMMARY tab back out — mirrors report.py's
    _load_sync_summary() convention (report reads the already-built sync Excel
    by path rather than recomputing). Returns (lga_rows, time_stats, note)
    where lga_rows are the per-LGA HIGH/MODERATE/LOW/NEVER SYNCED rows
    (same layout as cdd_sync.py's), time_stats = {label: (count, pct_str)}, and
    note is the disclosure text from the SUMMARY tab's note row (roster is
    sync-derived — see cdd_sync_itn.py docstring).
    Returns ([], {}, "") if the Excel doesn't exist yet (report_itn.py can run
    standalone before run.py's dispatch wiring calls cdd_sync_itn first).
    """
    if not path or not os.path.exists(path):
        return [], {}, ""
    try:
        wb = openpyxl.load_workbook(path, read_only=True)
        if "SUMMARY" not in wb.sheetnames:
            wb.close()
            return [], {}, ""
        ws = wb["SUMMARY"]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        wb.close()
    except Exception as e:
        log.warning(f"[report_itn] sync summary read failed (non-fatal): {e}")
        return [], {}, ""

    lga_rows, time_stats, note = [], {}, ""
    for r in rows:
        if not r or (not r[0] and not r[1]):
            continue
        label = str(r[0] or "").strip()
        if "Synced by" in label:
            time_stats[label] = (r[1], r[2])
        elif label and not r[1] and len(str(label)) > 40:
            note = label   # the merged note row has no LGA/count values
        elif r[1] and str(r[1]).strip().upper() != "GRAND TOTAL":
            # NOTE: strip the leading "#" column here (r[1:], not r) — report.py's
            # own _sync_table reads its equivalent rows with a `row[ci-1]` shift
            # that does NOT drop that leading column, so every field after LGA
            # renders one column off (LGA name appears under "Total CDDs", etc).
            # That's a pre-existing bug in report.py/cdd_sync.py we're not touching,
            # but deliberately not reproducing here — _sync_section_itn expects
            # row[0] to already be the LGA name.
            lga_rows.append(r[1:])
    return lga_rows, time_stats, note


# ── day-wise progress (ITN's equivalent of SPAQ's Days-1-to-N history) ─────────
# Mirrors report.py's _load_all_days_perf / _generate_progress_chart exactly now:
# analyze_itn.py applies a real GTE/LTE date filter (see its _date_filter), so
# each itn_history/performance_day{N}.xlsx snapshot is genuinely THAT DAY's own
# data — same as SPAQ's per-day files — and this accumulates a running sum across
# days, the same way report.py's _load_all_days_perf/_generate_progress_chart do,
# rather than needing any delta/subtraction workaround.

def _load_days_from_es_itn(cfg, elapsed_day):
    """
    Per-day series straight from ES AT REPORT TIME — one aggregation per
    campaign day (distinct households, sum of nets, sum of memberCount).
    This is the primary source for the cumulative view; the itn_history
    file-sum (_load_all_days_perf_itn) is only the fallback.

    Why ES and not the day files: each day file is frozen at that day's report
    time, and records synced later (with taskDates still = that day) never
    reach any file — on chad roughly half of a day's data arrives after the
    report runs, so the file-summed cumulative drifted to ~half of the
    dashboard's truth. ES always holds the complete picture for past days, so
    this back-fill self-heals on every run. Same pattern as report.py's
    _load_days_from_es for SPAQ.

    Semantics match the DSS dashboard tiles: households = distinct householdId
    per day (cardinality, approximate above ~40k/day), nets = sum(quantity),
    population = sum(memberCount) over ADMINISTRATION_SUCCESS records of the
    campaign. Raises on ES failure — the caller falls back to the file-sum.
    """
    import requests
    if not (cfg.get("campaign_start") and elapsed_day):
        return []
    url, idx, auth = cfg["es_url"], cfg["ES_INDEX_TASK"], cfg["es_auth"]
    date_field = cfg.get("task_date_field", "taskDates")
    scope = [
        {"term": {"Data.additionalDetails.projectReferenceId.keyword": cfg["campaign_number"]}},
        {"term": {"Data.administrationStatus.keyword": "ADMINISTRATION_SUCCESS"}},
    ]
    days = []
    cum_hh = cum_pop = cum_nets = 0
    for day_num in range(1, elapsed_day + 1):
        d   = cfg["campaign_start"] + timedelta(days=day_num - 1)
        rng = {"range": {f"Data.{date_field}": {
            "gte": f"{d.isoformat()}T00:00:00.000Z", "lte": f"{d.isoformat()}T23:59:59.999Z"}}}
        q = {"size": 0,
             "query": {"bool": {"filter": scope + [rng]}},
             "aggs": {
                 "hh":   {"cardinality": {"field": "Data.householdId.keyword",
                                           "precision_threshold": 40000}},
                 "nets": {"sum": {"field": "Data.quantity"}},
                 "pop":  {"sum": {"field": "Data.memberCount"}},
             }}
        r = requests.post(f"{url}/{idx}/_search", json=q, auth=auth, verify=False, timeout=60)
        r.raise_for_status()
        a = r.json()["aggregations"]
        hh, nets, pop = int(a["hh"]["value"]), int(a["nets"]["value"] or 0), int(a["pop"]["value"] or 0)
        cum_hh += hh; cum_pop += pop; cum_nets += nets
        days.append({
            "day": day_num, "date": d.strftime("%d %b"),
            "hh_visited": hh, "pop_covered": pop, "nets_distributed": nets,
            "cum_hh_visited": cum_hh, "cum_pop_covered": cum_pop,
            "cum_nets_distributed": cum_nets,
        })
    log.info(f"[report_itn] day series from ES: {len(days)} days, "
             f"cumulative nets {cum_nets:,} (late syncs included)")
    return days


def _load_all_days_perf_itn(cfg, elapsed_day):
    """
    Reads back each day's own nets_distributed/hh_visited/pop_covered ACTIVITY
    figures from itn_history/ and builds a running cumulative sum — these are
    genuine historical data (net counts don't go stale) so they're read as-is.

    Deliberately does NOT read/return each day's own stored target columns
    (Target HH/Pop/ITN) at all — a historical snapshot's target value can be
    stale (written by an older code version, before targets were divided by
    campaign length or before the zero-activity-LGA fix), and that staleness
    would silently leak into anything that used it. Every caller that needs a
    target uses today's current one instead (passed in separately) — see
    _generate_progress_chart_itn's and _progress_table_itn's docstrings.
    """
    days = []
    hist_dir = os.path.join(os.path.dirname(cfg["perf_xlsx"]), "itn_history")
    cum_hh_visited = cum_pop_covered = cum_nets_distributed = 0
    for day_num in range(1, (elapsed_day or 0) + 1):
        path = os.path.join(hist_dir, f"performance_day{day_num}.xlsx")
        if not os.path.exists(path):
            continue
        try:
            wb = openpyxl.load_workbook(path, read_only=True)
            ws = wb["ALL LGAS"]
            rows_raw = [
                r for r in ws.iter_rows(min_row=3, values_only=True)
                if r[2] and str(r[2]).strip() not in ("", "GRAND TOTAL")
            ]
            wb.close()
            # col order per analyze_itn.HEADERS: #, Province, LGA, Facilities,
            # Target HH(4), HH Visited(5), ..., Target Pop(7), Pop Covered(8), ...,
            # Target ITN(10), Nets Distributed(11), ... — only the *Visited/
            # *Covered/Distributed activity columns are read; the Target
            # columns are intentionally skipped (see docstring above).
            hh_visited       = sum(int(r[5] or 0)  for r in rows_raw)
            pop_covered      = sum(int(r[8] or 0)  for r in rows_raw)
            nets_distributed = sum(int(r[11] or 0) for r in rows_raw)

            cum_hh_visited        += hh_visited
            cum_pop_covered       += pop_covered
            cum_nets_distributed  += nets_distributed

            date_label = (cfg["campaign_start"] + __import__("datetime").timedelta(days=day_num - 1)).strftime("%d %b")
            days.append({
                "day": day_num, "date": date_label,
                "hh_visited": hh_visited, "pop_covered": pop_covered,
                "nets_distributed": nets_distributed,
                "cum_hh_visited": cum_hh_visited, "cum_pop_covered": cum_pop_covered,
                "cum_nets_distributed": cum_nets_distributed,
            })
        except Exception as e:
            log.warning(f"[report_itn] could not read itn_history/performance_day{day_num}.xlsx: {e}")
    return days


def _generate_progress_chart_itn(days_data, cfg, current_daily_target):
    """
    Bar chart, adapted from report.py's _generate_progress_chart. DELIBERATE
    DIFFERENCE from SPAQ's literal formula, disclosed: SPAQ's chart divides
    cumulative-treated by the FULL, fixed campaign target for every bar (so the
    chart only approaches 100% at the very end of the campaign) — but this
    module's Progress TABLE (see _progress_table_itn) shows a per-day PACING
    metric (cumulative distributed ÷ daily-target x elapsed-days-so-far), which
    is a materially different percentage. Showing two different formulas for
    what looks like "the same cumulative coverage" in the same section was
    confusing in practice, so the chart now uses the SAME pacing formula as the
    table — each bar's own denominator is current_daily_target x that bar's own
    elapsed day count, not one fixed full-campaign figure for every bar.

    current_daily_target is TODAY's freshly-loaded daily target, never a
    historical day's own stored value (see _load_all_days_perf_itn's docstring
    for why: a value written by pre-fix code can be permanently stale).
    """
    if not days_data:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        # Cap to the most recent 30 days so the chart stays legible on a
        # long-running campaign — full history still in itn_history/.
        shown = days_data[-30:]
        dropped = len(days_data) - len(shown)
        labels   = [f"Day {d['day']}\n{d['date']}" for d in shown]

        coverage = [
            (d["cum_nets_distributed"] / (current_daily_target * d["day"]) * 100
             if current_daily_target else 0)
            for d in shown
        ]

        def bar_color(c):
            if c >= 95: return "#1A7A1A"
            if c >= 70: return "#E06000"
            return "#CC0000"

        Y_MAX, BAR_CAP = 115, 110
        plot_vals = [min(c, BAR_CAP) for c in coverage]
        colors = [bar_color(c) for c in coverage]

        fig, ax = plt.subplots(figsize=(9, 4))
        fig.patch.set_facecolor("#F9F9F9")
        bars = ax.bar(labels, plot_vals, color=colors, width=0.5, zorder=3)
        title = f"{cfg['state_name']} — Cumulative ITN Coverage vs Campaign Target"
        if dropped > 0:
            title += f"  (last 30 of {len(days_data)} logged days)"
        ax.set_title(title, fontsize=11, fontweight="bold", pad=10)
        ax.set_ylabel("Cumulative ITN Coverage (%)", fontsize=9)
        ax.set_ylim(0, Y_MAX)
        ax.axhline(95, color="#1A7A1A", linestyle="--", linewidth=1, alpha=0.6, label="HIGH (95%)")
        ax.axhline(70, color="#E06000", linestyle="--", linewidth=1, alpha=0.6, label="MODERATE (70%)")
        ax.legend(fontsize=8)
        ax.grid(axis="y", linestyle="--", alpha=0.4, zorder=0)
        ax.set_axisbelow(True)
        for bar, val in zip(bars, coverage):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1.5,
                    f"{val:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold", clip_on=True)

        plt.tight_layout(pad=2.0)
        chart_path = os.path.join(os.path.dirname(cfg["perf_xlsx"]), f"itn_progress_chart_day{days_data[-1]['day']}.png")
        plt.savefig(chart_path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close()
        log.info(f"[report_itn] progress chart saved -> {chart_path}")
        return chart_path
    except Exception as e:
        log.warning(f"[report_itn] progress chart generation failed (non-fatal): {e}")
        return None


def _progress_table_itn(doc, days_data, current_daily_target):
    """
    Same shape as report.py's Section 4 table: each row is that DAY's own
    figures (not cumulative), with a cumulative TOTAL row at the bottom.

    report.py's own TOTAL row uses a DIFFERENT target base than its chart:
    cum_target = sum(d["target"] for d in days_data) — i.e. each elapsed day's
    OWN target summed, a "pacing" metric (are we hitting our daily targets on
    average so far), not the chart's full-campaign-target metric. Mirrored here
    as current_daily_target x elapsed days, rather than literally summing each
    day's stored target value — because a historical itn_history/ snapshot can
    carry a stale target (written by an older code version), and summing
    stale + fresh values would silently corrupt this total. Every day's real
    daily target is the same fixed value all campaign (full target / campaign
    days), so multiplying today's fresh figure by the day count is equivalent
    to the correct sum, without the stale-data risk.
    """
    if not days_data:
        add_para(doc, "No day-wise history yet — this is the first report run.")
        return
    shown = days_data[-14:]
    dropped = len(days_data) - len(shown)
    cols = ["Day", "Date", "Daily Target ITNs", "Nets Distributed", "Coverage"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for ci, h in enumerate(cols):
        hdr(table.cell(0, ci), h)
    # Every row uses TODAY's current daily target, NOT each day's own stored
    # net_target from its itn_history/ snapshot — the daily target is a fixed
    # value for the whole campaign (full target ÷ campaign days) that only ever
    # changes if the target file itself changes, so there is no need to trust a
    # historical snapshot's own copy of it. This is what actually fixes the
    # stale-Day-1 display, permanently, without needing to regenerate any old
    # file: a day written by pre-fix code showed the wrong (undivided) target in
    # its own snapshot, but that stored value is now simply never read for this.
    for ri, d in enumerate(shown, 1):
        row = table.add_row()
        alt = ri % 2 == 1
        day_cov_pct = (d["nets_distributed"] / current_daily_target * 100
                       if current_daily_target else 0)
        day_cov_str = f"{day_cov_pct:.1f}%" if current_daily_target else "N/A"
        cov_color = STATUS_COLOR.get(cov_band(day_cov_pct) if current_daily_target else "NO TARGET")
        vals = [f"Day {d['day']}", d["date"], f"{current_daily_target:,.0f}",
                f"{d['nets_distributed']:,}", day_cov_str]
        for ci, val in enumerate(vals):
            dat(row.cells[ci], val, alt=alt, color=cov_color if ci == 4 else None)

    # Cumulative totals row — current_daily_target x elapsed days (last["day"],
    # not len(days_data), so a gap doesn't under-count it), matching report.py's
    # cum_target = sum(each elapsed day's own target) without the stale-data risk.
    last = days_data[-1]
    cum_target = current_daily_target * last["day"]
    cum_cov_pct = last["cum_nets_distributed"] / cum_target * 100 if cum_target else 0
    tot_row = table.add_row()
    tot_vals = ["TOTAL", "", f"{cum_target:,}", f"{last['cum_nets_distributed']:,}",
                f"{cum_cov_pct:.1f}%" if cum_target else "N/A"]
    for ci, val in enumerate(tot_vals):
        dat(tot_row.cells[ci], val, bold=True)

    if dropped > 0:
        add_para(doc, f"Showing the most recent 14 of {len(days_data)} logged days. "
                       f"Full history in itn_history/.", size=8, color=GREY_RGB)


# ── narrative prompts (Groq) ─────────────────────────────────────────────────────

def _conclusion_prompt(cfg, g, hh_cov, pop_cov, net_cov, lga_d, roster_ever_synced=0, roster_high=0):
    best_lga  = max(lga_d, key=lambda d: lga_d[d]["nets_distributed"] / lga_d[d]["net_target"] * 100
                     if lga_d[d]["net_target"] else 0, default="N/A")
    worst_lga = min(lga_d, key=lambda d: lga_d[d]["nets_distributed"] / lga_d[d]["net_target"] * 100
                     if lga_d[d]["net_target"] else 100, default="N/A")
    sync_str = f"{roster_high:,}/{roster_ever_synced:,}" if roster_ever_synced else "N/A"
    header = f"{cfg['campaign_name']} {cfg['state_name']} ITN Distribution — Report as of {datetime.now().strftime('%Y-%m-%d')}"
    return (
        f"{header}\n"
        f"Household Coverage:{hh_cov} Population Coverage:{pop_cov} ITN Coverage:{net_cov}\n"
        f"Households Visited:{g['hh_visited']:,}/{g['hh_target']:,}  "
        f"Nets Distributed:{g['nets_distributed']:,}/{g['net_target']:,}\n"
        f"BestLGA:{best_lga} WorstLGA:{worst_lga}\n"
        f"MissingHHHead:{g['missing_hh_head']:,} MissingGPS:{g['missing_gps']:,}\n"
        f"CDDSync:{sync_str}\n"
        "Write a 5-sentence formal conclusion covering: (1) overall ITN distribution coverage vs target "
        "(2) best-performing LGA numbers (3) worst-performing LGA implication "
        "(4) data quality action needed (5) outlook for remaining distribution and CDD sync activity. Plain text only."
    )


def _slack_prompt(cfg, g, cum_hh, cum_pop, cum_nets, campaign_len, elapsed_day,
                  roster_ever_synced=0, roster_high=0):
    """Coverage in the Slack message is CUMULATIVE only:
    Cumulative Coverage % = Cumulative (Days 1-N) / Total Campaign Target x 100.
    No daily target and no daily coverage are exposed to the narrative."""
    sync_pct = f"{roster_high/roster_ever_synced*100:.1f}%" if roster_ever_synced else "N/A"
    n = campaign_len or 1
    tot_hh   = g["hh_target"]  if cfg.get("cumulative") else g["hh_target"]  * n
    tot_pop  = g["pop_target"] if cfg.get("cumulative") else g["pop_target"] * n
    tot_net  = g["net_target"] if cfg.get("cumulative") else g["net_target"] * n
    day_lbl  = elapsed_day or cfg.get("DAY", "?")
    facts = (
        f"Campaign: {cfg['campaign_name']} in {cfg['state_name']}, LLIN bed-net distribution, "
        f"Day {day_lbl} of {campaign_len or '?'}.\n"
        f"Cumulative ITN coverage (Days 1-{day_lbl} vs total campaign target): "
        f"{_cov_str(cum_nets, tot_net)}\n"
        f"Cumulative nets distributed (Days 1-{day_lbl}): {cum_nets:,} of {tot_net:,} total target\n"
        f"Cumulative household coverage: {_cov_str(cum_hh, tot_hh)} "
        f"({cum_hh:,} of {tot_hh:,})\n"
        f"Cumulative population coverage: {_cov_str(cum_pop, tot_pop)} "
        f"({cum_pop:,} of {tot_pop:,})\n"
        f"Households visited this day: {g['hh_visited']:,}\n"
        f"Nets distributed this day: {g['nets_distributed']:,}\n"
        # No duplicate figure: the legacy within-day dup_records number reads
        # as "double distribution" but only counts same-ID re-syncs — the
        # authoritative duplicate analysis lives in the internal report's
        # Duplicate Distribution section.
        f"Missing household head name: {g['missing_hh_head']:,}\n"
        f"Missing GPS: {g['missing_gps']:,}\n"
        + (f"CDDs synced: {roster_high:,} of {roster_ever_synced:,} ({sync_pct})\n" if roster_ever_synced else "")
    )
    return (
        facts
        + "\nWrite a single flowing paragraph (4 to 6 sentences) for a Slack field update "
          "summarising the figures above. Coverage must be described ONLY as cumulative "
          "coverage against the total campaign target; never mention a daily target or "
          "daily coverage. End with the single most urgent action for supervisors. "
          "Use the exact numbers given, do not invent any. Plain prose only — "
          "one paragraph, no heading, no bullet points, no emojis, no asterisks."
    )


def _issues_prompt(cfg, g, hh_cov, pop_cov, net_cov, lga_d, facilities, sync_lga_rows, sync_time_stats):
    # Deliberately receives NO duplicate-matrix facts: issues_data is shared
    # by the internal AND partner documents, and duplicate metrics are a DQ
    # surface stripped from every partner output. The matrix has its own
    # deterministic section in the internal report.
    low_act = [f for f in facilities if f["records"] < 10]
    worst_lga = min(
        lga_d,
        key=lambda d: lga_d[d]["nets_distributed"] / lga_d[d]["net_target"] * 100
        if lga_d[d]["net_target"] else 100,
        default="N/A",
    )
    total_cdds = sum(int(r[2] or 0) for r in sync_lga_rows if len(r) > 2 and r[2])
    dormant    = sum(int(r[3] or 0) for r in sync_lga_rows if len(r) > 3 and r[3])
    sync_str   = f"{total_cdds - dormant:,}/{total_cdds:,} active" if total_cdds else "N/A"
    by17 = ""
    for label, (count, pct) in (sync_time_stats or {}).items():
        if count is not None:
            by17 = f" by17:{count:,}({pct})"
    header = f"{cfg['campaign_name']} {cfg['state_name']} ITN Distribution — Report as of {datetime.now().strftime('%Y-%m-%d')}"
    return (
        f"{header}\n"
        f"HHCoverage:{hh_cov} PopCoverage:{pop_cov} ITNCoverage:{net_cov}\n"
        f"WorstLGA:{worst_lga}\n"
        f"LowActivity(<10 records):{len(low_act)} "
        f"MissingHHHead:{g['missing_hh_head']:,} MissingGPS:{g['missing_gps']:,}\n"
        f"CDDSync:{sync_str}{by17}\n"
        'Return ONLY a JSON array of 3-5 issues, no markdown:\n'
        '[{"observation":"...with numbers+LGA","status":"ACTIVE|RESOLVED","priority":"High|Moderate|Low",'
        '"notes":"actionable field instruction","data_type":"perf|sync"}]'
    )


def _generate_issues(cfg, g, hh_cov, pop_cov, net_cov, lga_d, facilities, sync_lga_rows, sync_time_stats):
    prompt = _issues_prompt(cfg, g, hh_cov, pop_cov, net_cov, lga_d, facilities, sync_lga_rows, sync_time_stats)
    raw = generate_narrative(prompt, max_tokens=1500)
    try:
        text = raw.strip()
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:])
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3]
        return json.loads(text)
    except Exception as e:
        try:
            t = text if text.startswith("[") else text[text.find("["):]
            cut = t.rfind("}")
            if cut != -1:
                salvaged = json.loads(t[:cut + 1].rstrip().rstrip(",") + "]")
                if isinstance(salvaged, list) and salvaged:
                    log.warning(f"[report_itn] issues JSON truncated ({e}) — salvaged {len(salvaged)} complete issue(s)")
                    return salvaged
        except Exception:
            pass
        log.warning(f"[report_itn] issues JSON parse failed: {e} — using raw text as single issue")
        return [{
            "observation": raw[:300] if raw else "No issues generated.",
            "status": "ACTIVE",
            "priority": "High",
            "notes": "Review manually.",
            "data_type": "perf",
        }]


# ── tables ─────────────────────────────────────────────────────────────────────

def _perf_table(doc, lga_d):
    header = ["LGA", "Facilities", "Target HH", "HH Visited", "HH Cov %",
              "Target ITN", "Nets Distributed", "ITN Cov %", "Status"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (dist, D) in enumerate(sorted(lga_d.items()), 1):
        raw_pct = D["nets_distributed"] / D["net_target"] * 100 if D["net_target"] else None
        cov     = f"{raw_pct:.1f}%" if raw_pct is not None else "N/A"
        stat    = cov_band(raw_pct) if raw_pct is not None else "NO TARGET"
        vals = [dist, D["facs"], f"{D['hh_target']:,}", f"{D['hh_visited']:,}",
                f"{D['hh_visited']/D['hh_target']*100:.1f}%" if D["hh_target"] else "N/A",
                f"{D['net_target']:,}", f"{D['nets_distributed']:,}", cov, stat]
        row = table.add_row()
        alt = ri % 2 == 1
        for ci, val in enumerate(vals):
            if alt:
                set_cell_bg(row.cells[ci], "EBF3FB")
            set_cell_borders(row.cells[ci])
            p = row.cells[ci].paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.name = FONT; run.font.size = Pt(9)
            if ci in (7, 8):
                run.bold = True
                color = STATUS_COLOR.get(stat)
                if color: run.font.color.rgb = color
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER


def _dq_table_itn(doc, lga_d):
    """Per-LGA DQ breakdown — mirrors report.py's _dq_table (3.2 in SPAQ)."""
    header = ["LGA", "Duplicates", "Missing HH Head", "Missing GPS", "Manual Codes", "Missing Codes"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (dist, D) in enumerate(sorted(lga_d.items()), 1):
        vals = [dist, D["dup_records"], D["missing_hh_head"], D["missing_gps"],
                D["manual_codes"], D["missing_codes"]]
        row = table.add_row()
        alt = ri % 2 == 1
        for ci, val in enumerate(vals):
            dat(row.cells[ci], val, alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)


def _sync_section_itn(doc, sec_num, sync_lga_rows, sync_time_stats, sync_note,
                       roster_ever_synced, roster_high, sync_link=""):
    """
    CDD Sync Activity — ITN's equivalent of SPAQ's Section 5. Same HIGH/MODERATE/
    LOW/NEVER SYNCED status model and column layout as cdd_sync.py's Section 5.1/
    5.1b. ONE disclosed structural difference: the roster is derived from sync
    records themselves (no assigned-staff roster available — see cdd_sync_itn.py's
    docstring), so a CDD with zero syncs EVER cannot appear; NEVER SYNCED counts
    only roster CDDs whose syncs all predate campaign_start.
    """
    add_heading(doc, f"{sec_num}.1  FLW Sync Summary", 5)
    if not roster_ever_synced:
        add_para(doc, "CDD sync data not available for this report (cdd_sync_itn.py "
                       "may not have run yet).", size=9, color=GREY_RGB)
        return
    sync_pct = f"{roster_high/roster_ever_synced*100:.1f}%" if roster_ever_synced else "N/A"
    # NEVER SYNCED from the SUMMARY rows (col 5 after the # strip) — no longer a
    # hardcoded 0: window-scoping means pre-campaign-only CDDs now count as NEVER
    never_total = 0
    for r in (sync_lga_rows or []):
        try:
            never_total += int(r[5] or 0)
        except (ValueError, TypeError, IndexError):
            pass
    summary_rows = [
        ("Total CDDs Registered",   f"{roster_ever_synced:,}"),
        ("CDDs Synced (total day)", f"{roster_high:,}  ({sync_pct})"),
        ("Never Synced",            f"{never_total:,}"),
    ]
    for label, (count, pct) in (sync_time_stats or {}).items():
        hour = label.replace("Synced by ", "").replace(" today (UTC)", "")
        if count is not None:
            summary_rows.append((f"Synced by {hour}", f"{count:,}  ({pct})  — early sync indicator"))
    two_col_table(doc, summary_rows, col_widths=(5, 6))
    doc.add_paragraph()

    if sync_note:
        add_para(doc, sync_note, size=8, color=GREY_RGB)
        doc.add_paragraph()

    if sync_lga_rows:
        add_heading(doc, f"{sec_num}.2  Sync Status by LGA", 5)
        note_p = add_para(doc, "", size=9, color=GREY_RGB)
        if sync_link:
            add_hyperlink(note_p, "Full CDD sync data ↗", sync_link)
        else:
            run_d = note_p.add_run("Full CDD sync data.")
            run_d.font.name = FONT; run_d.font.size = Pt(9); run_d.font.color.rgb = GREY_RGB
        cols = ["#", "LGA", "Total CDDs", "HIGH", "MODERATE", "LOW", "NEVER SYNCED", "% Never Synced"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for ci, h in enumerate(cols):
            hdr(table.cell(0, ci), h)
        for ri, row in enumerate(sync_lga_rows, 1):
            tr = table.add_row()
            alt = ri % 2 == 1
            for ci in range(len(cols)):
                val = ri if ci == 0 else (row[ci - 1] if ci - 1 < len(row) else "")
                dat(tr.cells[ci], val, alt=alt)


def _dq_summary_table(doc, g, sec_num="3.5", dup_matrix=None, lga_d=None, perf_link=""):
    total = g["hh_visited"] or 1
    # The flat "Duplicate Records" row appears only when the matrix wasn't
    # measured — otherwise its four-way breakdown (subsection .3) replaces it.
    metrics = []
    if dup_matrix is None:
        metrics.append(("Duplicate Records", g["dup_records"]))
    metrics += [
        ("Missing Household Head Name",    g["missing_hh_head"]),
        ("Missing GPS",                    g["missing_gps"]),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr(table.cell(0, 0), "Metric")
    hdr(table.cell(0, 1), "Count")
    hdr(table.cell(0, 2), "% of Households Visited")
    for ri, (metric, count) in enumerate(metrics, 1):
        row = table.add_row()
        pct = f"{count/total*100:.2f}%"
        dat(row.cells[0], metric, alt=ri % 2 == 1, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f"{count:,}", alt=ri % 2 == 1)
        dat(row.cells[2], pct, alt=ri % 2 == 1)
    doc.add_paragraph()

    # ITN-specific addition, no SPAQ/AZM equivalent — mirrors the campaign dashboard's own
    # headline DQ metric: manual code entry is far more error/fraud-prone than barcode scanning.
    total_codes = g["manual_codes"] + g["scanned_codes"]
    pct_scanned = f"{g['scanned_codes']/total_codes*100:.2f}%" if total_codes else "N/A"
    pct_manual  = f"{g['manual_codes']/total_codes*100:.2f}%" if total_codes else "N/A"
    add_heading(doc, f"{sec_num}.1  Bednet Code Entry Method", 5)
    code_table = doc.add_table(rows=1, cols=3)
    code_table.style = "Table Grid"
    hdr(code_table.cell(0, 0), "Entry Method")
    hdr(code_table.cell(0, 1), "Count")
    hdr(code_table.cell(0, 2), "% of Codes")
    for ri, (label, count, pct) in enumerate([
        ("Scanned", g["scanned_codes"], pct_scanned),
        ("Manual",  g["manual_codes"],  pct_manual),
    ], 1):
        row = code_table.add_row()
        dat(row.cells[0], label, alt=ri % 2 == 1, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f"{count:,}", alt=ri % 2 == 1)
        dat(row.cells[2], pct, alt=ri % 2 == 1)
    doc.add_paragraph()

    # Missing Codes — records where a net was distributed but NEITHER scanned nor
    # manually recorded (zero barcode documentation, untraceable in inventory).
    # Distinct from the scanned-vs-manual RATIO above; denominator is total
    # records, not total codes (a missing-codes record contributes 0 to either).
    total_records = g["records"] or 1
    pct_missing_codes = f"{g['missing_codes']/total_records*100:.2f}%"
    add_heading(doc, f"{sec_num}.2  Missing Bednet Codes", 5)
    miss_table = doc.add_table(rows=1, cols=3)
    miss_table.style = "Table Grid"
    hdr(miss_table.cell(0, 0), "Metric")
    hdr(miss_table.cell(0, 1), "Count")
    hdr(miss_table.cell(0, 2), "% of Records")
    row = miss_table.add_row()
    dat(row.cells[0], "Neither scanned nor manually entered", align=WD_ALIGN_PARAGRAPH.LEFT)
    dat(row.cells[1], f"{g['missing_codes']:,}")
    dat(row.cells[2], pct_missing_codes)
    doc.add_paragraph()

    # Duplicate Distribution — rendered only when the matrix was measured.
    if dup_matrix is not None:
        add_heading(doc, f"{sec_num}.3  Duplicate Distribution", 5)
        _dup_matrix_section(doc, dup_matrix, lga_d or {}, g["records"], perf_link=perf_link)
        doc.add_paragraph()


def _dup_matrix_section(doc, dm, lga_d, total_records, perf_link=""):
    """
    Duplicate Distribution table — internal report only (the caller gates on
    `partner`, like every DQ section). Counts are record-level classifications
    against each household's earliest campaign-to-date record (see
    analyze_itn._classify_duplicates), so they deliberately do not reconcile
    with the within-day Duplicate Records column — the note under the table
    says so.
    """
    total_records = total_records or 1
    rows_spec = [
        ("Same user, same day", dm["dup_su_sd"],
         "Sync retries / double taps — technical, no field action", False),
        ("Same user, different days", dm["dup_su_dd"],
         "CDD re-visited a household they already served", False),
        ("Different users, same day", dm["dup_du_sd"],
         "Two CDDs served one household the same day — coordination gap", True),
        ("Different users, different days", dm["dup_du_dd"],
         "Household re-served on a later day — double distribution", True),
    ]

    # Equation only — the classification logic lives in the grey note below.
    add_para(doc, "Formula:  % of Records = Bucket Count ÷ Total Records × 100",
              size=8, color=GREY_RGB)

    _DUP_RED = RGBColor(0xCC, 0x00, 0x00)   # same red as ACTIVE/High elsewhere
    header = ["Duplicate Type", "Count", "% of Records", "Reading"]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (label, count, reading, emph) in enumerate(rows_spec, 1):
        row = table.add_row()
        alt = ri % 2 == 1
        # The costly buckets (different-user / re-served rows) flag in bold red,
        # and only when non-zero — a zero is good news, not an alert.
        flag = _DUP_RED if (emph and count) else None
        dat(row.cells[0], label, alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT, bold=emph)
        dat(row.cells[1], f"{count:,}", alt=alt, bold=emph, color=flag)
        dat(row.cells[2], f"{count/total_records*100:.2f}%", alt=alt, bold=emph, color=flag)
        dat(row.cells[3], reading, alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
    doc.add_paragraph()

    # Method note, ending in a Drive hyperlink to the performance Excel (same
    # pattern as the Low Activity and sync sections' links).
    note_p = add_para(doc, "Households are matched by head name, household size and village; "
                            "any delivery after a household's first is a duplicate. Affected "
                            "households are listed per type in the DUP sheets of the ",
                       size=8, color=GREY_RGB)
    if perf_link:
        add_hyperlink(note_p, "performance Excel ↗", perf_link)
    else:
        run = note_p.add_run("performance Excel.")
        run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = GREY_RGB


def _low_activity_facility_table(doc, facilities, perf_link=""):
    """
    Low Activity Facilities — records < 10. ADAPTED from SPAQ/AZM's LOW-coverage-
    PLUS-Low-Activity table: a facility-level Target/Coverage % no longer exists
    (targets only exist at LGA grain now — see analyze_itn.py's docstring), so
    this can only rank by raw activity (records), not coverage against a target.
    """
    low_facs = [f for f in facilities if f["records"] < 10]
    low_facs = sorted(low_facs, key=lambda f: f["records"])

    total_facs = len(facilities)
    shown      = min(len(low_facs), 20)
    note_text  = f"Showing {shown} Low Activity facilities (<10 records) of {total_facs} total. "
    note_p     = add_para(doc, note_text, size=9, color=GREY_RGB)
    if perf_link:
        add_hyperlink(note_p, "Full list in performance Excel ↗", perf_link)
    else:
        run = note_p.add_run("Full list in performance Excel.")
        run.font.name = FONT; run.font.size = Pt(9); run.font.color.rgb = GREY_RGB

    if not low_facs:
        add_para(doc, "No Low Activity facilities.")
        return
    header = ["#", "Province", "LGA", "Health Facility", "Records", "Households Visited", "Nets Distributed"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, f in enumerate(low_facs[:20], 1):
        row = table.add_row()
        alt = ri % 2 == 1
        vals = [ri, f["province"], f["lga"], f["fac"], f["records"],
                f"{f['households_visited']:,}", f"{f['nets_distributed']:,}"]
        for ci, val in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if ci in (1, 2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            dat(row.cells[ci], val, alt=alt, align=align)


# ── partner-safe Excel (DQ columns stripped) ────────────────────────────────────

# Same DQ-column set as analyze_itn.HEADERS' DQ additions — kept in sync manually
# since this reads the Excel by header name, not by importing analyze_itn.
_DQ_COLS_ITN = {"Duplicate Records", "Missing HH Head", "Missing GPS",
                "Manual Codes", "Scanned Codes", "% Scanned", "Missing Codes",
                "Dup Same User Same Day", "Dup Same User Diff Day",
                "Dup Diff User Same Day", "Dup Diff User Diff Day"}


def _make_partner_perf_xlsx_itn(perf_path):
    """
    Write a partner-safe copy of the ITN performance Excel with DQ columns stripped
    from every tab — mirrors report.py's _make_partner_perf_xlsx() for SPAQ/AZM.
    """
    from openpyxl.utils import get_column_letter
    HEADER_ROW = 2   # row 1 = banner, row 2 = column headers
    if not perf_path or not os.path.exists(perf_path):
        return ""
    try:
        wb = openpyxl.load_workbook(perf_path)
        # Every DUP* sheet (one per duplicate type; also covers the older
        # single "DUPLICATE DETAIL" tab) is dropped WHOLE, not column-stripped —
        # head names + household CRIDs + usernames, internal tracing data that
        # must never leave the internal report.
        for name in [n for n in wb.sheetnames if n.upper().startswith("DUP")]:
            wb.remove(wb[name])
        for ws in wb.worksheets:
            headers = [c.value for c in ws[HEADER_ROW]]
            to_del  = sorted([i + 1 for i, h in enumerate(headers) if h in _DQ_COLS_ITN], reverse=True)
            if not to_del:
                continue
            merges = [str(m) for m in list(ws.merged_cells.ranges) if m.min_row == 1]
            for m in merges:
                ws.unmerge_cells(m)
            for ci in to_del:
                ws.delete_cols(ci, 1)
            if merges:
                ws.merge_cells(f"A1:{get_column_letter(ws.max_column)}1")
        base, ext = os.path.splitext(perf_path)
        out = f"{base}_partner{ext}"
        wb.save(out)
        log.info(f"[report_itn] partner performance Excel written -> {out}")
        return out
    except Exception as e:
        log.warning(f"[report_itn] partner perf Excel build failed (non-fatal): {e}")
        return ""


# ── document builder ───────────────────────────────────────────────────────────

def _build_doc(cfg, *, g, hh_cov, pop_cov, net_cov, lga_d, facilities,
                issues_data, sync_lga_rows, sync_time_stats, sync_note,
                roster_ever_synced, roster_high, days_data, chart_path,
                conclusion, perf_link, sync_link="", perf_path="", sync_path="",
                partner=False, dup_matrix=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.8)
        section.left_margin = section.right_margin = Cm(2.0)

    # Elapsed day / campaign length — computed once here, reused by both the
    # header block below and the Daily Target section further down. Sourced from
    # cfg["campaign_start"]/["campaign_end"]/["extract_date"] (all guaranteed
    # present — config.py requires campaign_start/campaign_end on every row),
    # never from cfg["DAY"]/["campaign_days"] (config.py defaults campaign_days
    # to 4 when a sheet row doesn't set it, which would silently misreport the
    # day count here).
    campaign_start = cfg.get("campaign_start")
    campaign_end   = cfg.get("campaign_end")
    extract_date   = cfg.get("extract_date")
    campaign_days_total = (
        (campaign_end - campaign_start).days + 1
        if campaign_start and campaign_end else None
    )
    elapsed_day = (
        (extract_date - campaign_start).days + 1
        if campaign_start and extract_date else None
    )
    is_day_one = elapsed_day == 1

    title_p = doc.add_paragraph(style="Normal")
    title_p.clear()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"{cfg['state_name']}  —  {cfg['campaign_name']}")
    run.font.name = FONT; run.font.size = Pt(26)
    run.bold = True; run.font.color.rgb = TITLE_RGB

    # Subtitle — same "campaign dates — Day N of M" framing as report.py's
    # non-cumulative subtitle (Oyo's "30 July 2026 to 2 August 2026 — Day 2 of 4").
    sub = doc.add_paragraph(style="Normal")
    sub.clear()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_text = (
        f"{cfg['START_LABEL']} to {cfg['END_LABEL']}  —  Day {elapsed_day} of {campaign_days_total}"
        if elapsed_day and campaign_days_total else
        f"{cfg['START_LABEL']} to {cfg['END_LABEL']}"
    )
    r2 = sub.add_run(sub_text)
    r2.font.name = FONT; r2.font.size = Pt(14)
    r2.bold = True; r2.font.color.rgb = TITLE_RGB

    # Grey summary line — same content/style as report.py's (Day N — Date |
    # Extract timestamp | LGA count | Coverage), ITN Coverage as the headline
    # since that's the campaign's primary metric (nets distributed vs target).
    day_label = (extract_date.strftime("%d %B %Y") if extract_date
                 else datetime.now().strftime("%d %B %Y"))
    period = (f"Day {elapsed_day}  —  {day_label}" if elapsed_day else day_label)
    summary_line = (
        f"{period}  |  "
        f"Extract: {day_label}, {datetime.now().strftime('%H:%M')}  |  "
        f"{len(lga_d)} LGAs  |  Coverage: {net_cov}"
    )
    grey_p = add_para(doc, summary_line, size=8, color=GREY_RGB)
    grey_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Section numbering is dynamic (not hardcoded) since the partner report skips
    # the DQ subsection and the section count differs — same idea as report.py's
    # fac_sec/nonadmin_sec computed strings, generalised across the whole doc.
    sec = 1

    cum = cfg.get("cumulative", False)
    sec1_title = "Cumulative Campaign Overview" if cum else "Day Operational Overview"
    add_heading(doc, f"{sec}.  {sec1_title}", 4)
    overview_rows = [
        ("State / Country",                 cfg["state_name"]),
        ("Activity",                        "LLIN Bed-Net Distribution"),
        ("Campaign Dates",                  f"{cfg['START_LABEL']} to {cfg['END_LABEL']}"),
        ("Data Extract Timestamp",          datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("LGAs Covered",                    f"{len(lga_d)}"),
    ]

    if cum:
        # cfg["GTE"]/["LTE"] span the whole campaign in this mode (set by
        # run_cumulative()) and analyze_itn.py doesn't divide targets, so g IS
        # already the full-campaign total — no daily split, matching report.py's
        # cumulative-mode overview_rows exactly (no "Daily X" rows at all).
        overview_rows += [
            ("Overall Target Households",   f"{g['hh_target']:,}"),
            ("Total Households Visited",    f"{g['hh_visited']:,}"),
            ("Household Coverage",          hh_cov),
            ("Overall Target Population",   f"{g['pop_target']:,}"),
            ("Total Population Covered",    f"{g['pop_covered']:,}"),
            ("Population Coverage",         pop_cov),
            ("Overall Target ITNs",         f"{g['net_target']:,}"),
            ("Total Nets Distributed",      f"{g['nets_distributed']:,}"),
            ("ITN Coverage",                net_cov),
        ]
    else:
        # Daily mode: g IS today's own data (analyze_itn.py applies a real
        # GTE/LTE date filter now — see its _date_filter), and its targets are
        # already divided to a daily figure — same "Total Campaign Target =
        # Daily Target x Campaign Days" relationship as report.py's daily mode.
        # Cumulative figures are the running sum across itn_history/ day-files
        # (days_data[-1]'s cum_* fields — see _load_all_days_perf_itn), exactly
        # like report.py sums multiple days' own files for its cumulative view.
        last = days_data[-1] if days_data else None
        cum_hh_visited       = last["cum_hh_visited"]       if last else g["hh_visited"]
        cum_pop_covered      = last["cum_pop_covered"]      if last else g["pop_covered"]
        cum_nets_distributed = last["cum_nets_distributed"] if last else g["nets_distributed"]

        # Cumulative block FIRST (headline, vs TOTAL campaign targets — same
        # basis as the Slack message), daily block second.
        _len = None
        if cfg.get("campaign_start") and cfg.get("campaign_end"):
            _len = (cfg["campaign_end"] - cfg["campaign_start"]).days + 1
        _n = _len or 1
        d_lbl = elapsed_day or '?'
        overview_rows += [
            ("Total Campaign Target ITNs",    f"{g['net_target'] * _n:,}"),
            (f"Cumulative Nets Distributed (Days 1-{d_lbl})",  f"{cum_nets_distributed:,}"),
            (f"Cumulative ITN Coverage (Days 1-{d_lbl})",
             _cov_str(cum_nets_distributed, g["net_target"] * _n)),
            ("Total Campaign Target Households", f"{g['hh_target'] * _n:,}"),
            (f"Cumulative Households Visited (Days 1-{d_lbl})", f"{cum_hh_visited:,}"),
            (f"Cumulative HH Coverage (Days 1-{d_lbl})",
             _cov_str(cum_hh_visited, g["hh_target"] * _n)),
            ("Total Campaign Target Population", f"{g['pop_target'] * _n:,}"),
            (f"Cumulative Population Covered (Days 1-{d_lbl})", f"{cum_pop_covered:,}"),
            (f"Cumulative Pop Coverage (Days 1-{d_lbl})",
             _cov_str(cum_pop_covered, g["pop_target"] * _n)),
            ("Daily Target Households",       f"{g['hh_target']:,}"),
            ("Households Visited Today",      f"{g['hh_visited']:,}"),
            ("Coverage vs Daily HH Target",   _cov_str(g["hh_visited"], g["hh_target"])),
            ("Daily Target Population",       f"{g['pop_target']:,}"),
            ("Population Covered Today",      f"{g['pop_covered']:,}"),
            ("Coverage vs Daily Pop Target",  _cov_str(g["pop_covered"], g["pop_target"])),
            ("Daily Target ITNs",             f"{g['net_target']:,}"),
            ("Nets Distributed Today",        f"{g['nets_distributed']:,}"),
            ("Coverage vs Daily ITN Target",  _cov_str(g["nets_distributed"], g["net_target"])),
        ]
    two_col_table(doc, overview_rows)
    doc.add_paragraph()

    # Coverage chart — directly under Section 1 overview table, same position as
    # report.py's chart (NOT inside Campaign Progress — that section is table-only).
    if chart_path and os.path.exists(chart_path):
        formula_p = add_para(
            doc,
            "Formula:  Cumulative ITN Coverage % = Cumulative Nets Distributed  ÷  "
            "Cumulative Target ITNs × 100",
            size=8, color=GREY_RGB,
        )
        formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(chart_path, width=Inches(5.5))
        doc.add_paragraph()
    sec += 1

    # Section 2 — Program Issues (Claude-generated, same model as report.py's
    # Section 2, adapted facts — see _issues_prompt above)
    add_heading(doc, f"{sec}.  Program Issues and Resolutions", 4)
    add_para(doc, "Issues identified across the extract and field reports.", size=9, color=GREY_RGB)
    issues_header = ["#", "Issue / Observation", "Status", "Priority", "Notes", "Data"]
    tbl2 = doc.add_table(rows=1, cols=len(issues_header))
    tbl2.style = "Table Grid"
    for ci, h in enumerate(issues_header):
        hdr(tbl2.cell(0, ci), h)
    from docx.shared import RGBColor
    _ACTIVE_RED   = RGBColor(0xCC, 0x00, 0x00)
    _RESOLVED_GRN = RGBColor(0x1A, 0x7A, 0x1A)
    _PRI_HIGH     = RGBColor(0xCC, 0x00, 0x00)
    _PRI_MOD      = RGBColor(0xE0, 0x60, 0x00)
    for ri, issue in enumerate(issues_data or [], 1):
        row2 = tbl2.add_row()
        alt = ri % 2 == 1
        status   = str(issue.get("status", "ACTIVE")).upper()
        priority = str(issue.get("priority", "High"))
        dtype    = str(issue.get("data_type", "perf")).lower()
        dat(row2.cells[0], ri, alt=alt)
        dat(row2.cells[1], issue.get("observation", ""), alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        dat(row2.cells[2], status, alt=alt, color=_ACTIVE_RED if status == "ACTIVE" else _RESOLVED_GRN)
        dat(row2.cells[3], priority, alt=alt,
            color=_PRI_HIGH if priority == "High" else _PRI_MOD if priority == "Moderate" else None)
        dat(row2.cells[4], issue.get("notes", ""), alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT, size=9)

        # Data column — hyperlink to relevant Drive file, same pattern as report.py.
        data_cell = row2.cells[5]
        if alt:
            set_cell_bg(data_cell, ALT_FILL)
        set_cell_borders(data_cell)
        dp = data_cell.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        link  = sync_link if dtype == "sync" else perf_link
        label = "CDD Sync Data" if dtype == "sync" else "Performance Data"
        fname = (os.path.basename(sync_path) if dtype == "sync" else os.path.basename(perf_path)) or label

        if link:
            add_hyperlink(dp, label + " ↗", link)
        else:
            run_d = dp.add_run(fname)
            run_d.font.name = FONT; run_d.font.size = Pt(8)
            run_d.font.color.rgb = GREY_RGB
    for ci, w in enumerate([Cm(0.8), Cm(5.5), Cm(1.5), Cm(1.6), Cm(4.5), Cm(2.5)]):
        for cell in tbl2.columns[ci].cells:
            cell.width = w
    doc.add_paragraph()
    sec += 1

    dist_sec = sec
    add_heading(doc, f"{dist_sec}.  Distribution Data Analysis", 4)
    if not cum:
        _fl = None
        if cfg.get("campaign_start") and cfg.get("campaign_end"):
            _fl = (cfg["campaign_end"] - cfg["campaign_start"]).days + 1
        if _fl:
            add_para(doc,
                     f"All targets and coverage in this section are DAILY: "
                     f"Daily Target = Total Campaign Target \u00f7 Campaign Length "
                     f"(ITNs: {g['net_target'] * _fl:,} \u00f7 {_fl} days "
                     f"= {g['net_target']:,} per day).",
                     size=9, color=GREY_RGB)
    sub = 1
    add_heading(doc, f"{dist_sec}.{sub}  Performance by LGA", 5); sub += 1
    _perf_table(doc, lga_d)
    doc.add_paragraph()

    if not partner:
        add_heading(doc, f"{dist_sec}.{sub}  Data Quality by LGA", 5); sub += 1
        _dq_table_itn(doc, lga_d)
        doc.add_paragraph()

    add_heading(doc, f"{dist_sec}.{sub}  Low Activity Facilities", 5); sub += 1
    add_para(doc, "Facilities with fewer than 10 records. Sorted ascending — no facility-level "
                   "target exists (targets are only available at LGA grain).", size=9, bold=True)
    _low_activity_facility_table(doc, facilities, perf_link=perf_link)
    doc.add_paragraph()

    # Non-Administration Analysis (SPAQ/AZM's equivalent section) intentionally
    # NOT included yet — the household-delivery equivalent of Absent/Refused/
    # Ineligible/Referred/Died/Migrated statuses for chad has not been verified
    # against live ES (only ADMINISTRATION_SUCCESS has ever been queried this
    # session). Do not fabricate zeros here; add once confirmed.

    if not partner:
        # Duplicate Distribution renders inside Data Quality Summary (as .3,
        # replacing the flat Duplicate Records row) and only when the matrix
        # was measured — None keeps the summary's classic shape.
        add_heading(doc, f"{dist_sec}.{sub}  Data Quality Summary", 5)
        _dq_summary_table(doc, g, sec_num=f"{dist_sec}.{sub}",
                          dup_matrix=dup_matrix, lga_d=lga_d, perf_link=perf_link)
        doc.add_paragraph()
    sec += 1

    # Campaign Progress — ITN's equivalent of SPAQ's Section 4. Reads the
    # itn_history/ day-file series analyze_itn.py now writes (see that module's
    # run() and _load_all_days_perf_itn's docstrings for the mechanism).
    add_heading(doc, f"{sec}.  Campaign Progress  —  {cfg['START_LABEL']} to {cfg['END_LABEL']}", 4)
    add_para(doc, "Coverage = Nets Distributed ÷ Daily Target ITNs. Cumulative figures "
                   "include all days to date. Chart above, in Section 1.",
              size=9, color=GREY_RGB)
    # Table needs the TRUE daily target regardless of report mode — in
    # cumulative mode g["net_target"] is already the full campaign figure
    # (analyze_itn.py doesn't divide it there), so normalise it back down.
    current_daily_target = (g["net_target"] / (campaign_days_total or 1)) if cum else g["net_target"]
    _progress_table_itn(doc, days_data, current_daily_target)
    doc.add_paragraph()
    sec += 1

    # CDD Sync Activity — ITN's equivalent of SPAQ's Section 5
    add_heading(doc, f"{sec}.  CDD Sync Activity", 4)
    _sync_section_itn(doc, sec, sync_lga_rows, sync_time_stats, sync_note,
                       roster_ever_synced, roster_high, sync_link=sync_link)
    doc.add_paragraph()
    sec += 1

    add_heading(doc, f"{sec}.  Conclusion", 4)
    add_para(doc, conclusion, size=10)

    from datetime import timezone
    extracted_at = datetime.now(timezone.utc).strftime("%H:%M UTC")
    disclaimer = (
        f"Data extracted from DIGIT HCM as at {extracted_at}. "
        f"Figures are subject to change as field teams synchronise."
    )
    section = doc.sections[0]
    section.footer_distance = Cm(0.8)
    footer_para = section.footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_para.add_run(disclaimer)
    fr.font.size = Pt(7); fr.font.color.rgb = GREY_RGB; fr.font.name = FONT

    return doc


# ── public entry point ─────────────────────────────────────────────────────────

def run(cfg):
    log.info(f"[report_itn] {cfg['state_name']} ...")
    perf_path = cfg["perf_xlsx"]
    if not os.path.exists(perf_path):
        raise FileNotFoundError(f"ITN performance Excel not found: {perf_path}")

    lga_d, facilities = _load_perf_itn(perf_path)
    g = _grand_totals(lga_d)
    hh_cov  = _cov_str(g["hh_visited"], g["hh_target"])
    pop_cov = _cov_str(g["pop_covered"], g["pop_target"])
    net_cov = _cov_str(g["nets_distributed"], g["net_target"])

    # Duplicate-distribution matrix — None when not measured for this extract
    # (dup_matrix off, enrichment failed, or a pre-matrix Excel); every consumer
    # below (doc section, issues facts, Slack line) skips entirely on None.
    dup_matrix = _dup_matrix_totals(lga_d)

    log.info(f"  {len(facilities)} facilities, {len(lga_d)} LGAs, ITN coverage {net_cov}")

    # CDD sync data — read back the Excel cdd_sync_itn.py already wrote (same
    # convention as report.py reading cfg["sync_xlsx"], not recomputing here).
    sync_path = cfg.get("sync_xlsx", "")
    sync_lga_rows, sync_time_stats, sync_note = _load_sync_summary_itn(sync_path)
    # Stripped row layout (see _load_sync_summary_itn): LGA, Total CDDs,
    # HIGH, MODERATE, LOW, NEVER SYNCED, % Never Synced.
    roster_ever_synced = sum(int(r[1] or 0) for r in sync_lga_rows if len(r) > 1 and r[1])
    roster_high        = sum(int(r[2] or 0) for r in sync_lga_rows if len(r) > 2 and r[2])
    if not sync_lga_rows:
        log.warning("  CDD sync summary not found/empty — run cdd_sync_itn.py first for a complete report")

    perf_link, sync_link = "", ""
    if not cfg.get("no_upload"):
        try:
            from dst_data_analysis_report.pipeline import notify as _notify
            fid   = _notify.campaign_folder_id(cfg)
            now_hm = datetime.now().strftime("%Y-%m-%d %H:%M")
            perf_title = f"{cfg['state_name']} ITN Performance Data — {now_hm}"
            perf_link = _notify.upload_file(perf_path, perf_title, folder_id=fid)
            if sync_path and os.path.exists(sync_path):
                sync_title = f"{cfg['state_name']} ITN CDD Sync Data — {now_hm}"
                sync_link = _notify.upload_file(sync_path, sync_title, folder_id=fid)
        except Exception as e:
            log.warning(f"  Drive upload failed (non-fatal): {e}")

    # Stash Drive links on cfg for the caller — mirrors report.py, used by
    # run.py's run_cumulative() summary printout (harmless for daily runs).
    cfg["perf_drive_link"] = perf_link
    cfg["sync_drive_link"] = sync_link

    # Day-wise progress — primary source is ES at report time (late syncs
    # included; the itn_history files are frozen at each day's report time and
    # undercount — proven ~2x low on chad). File-sum is the fallback only.
    elapsed_day = None
    if cfg.get("campaign_start") and cfg.get("extract_date"):
        elapsed_day = (cfg["extract_date"] - cfg["campaign_start"]).days + 1
    days_data = []
    try:
        days_data = _load_days_from_es_itn(cfg, elapsed_day)
    except Exception as e:
        log.warning(f"[report_itn] ES day-series back-fill failed (non-fatal — "
                    f"falling back to itn_history files, cumulative may undercount): {e}")
    if not days_data:
        days_data = _load_all_days_perf_itn(cfg, elapsed_day)

    # TRUE daily target, generic for any campaign/tenant: g["net_target"] is
    # already the daily figure in daily mode, or the full campaign figure in
    # cumulative mode (see analyze_itn.py's run()) — normalise the cumulative
    # case back down to a daily figure by dividing by the real campaign length
    # (cfg["campaign_start"]/["campaign_end"], required for every row by
    # config.py). Passed to the chart/table instead of letting them read any
    # historical day's own stored target — see _generate_progress_chart_itn's
    # docstring for why a stale historical value must never be used.
    campaign_days_for_target = None
    if cfg.get("campaign_start") and cfg.get("campaign_end"):
        campaign_days_for_target = (cfg["campaign_end"] - cfg["campaign_start"]).days + 1
    current_daily_target = (
        g["net_target"] / (campaign_days_for_target or 1) if cfg.get("cumulative")
        else g["net_target"]
    )
    chart_path = _generate_progress_chart_itn(days_data, cfg, current_daily_target)

    issues_data = _generate_issues(
        cfg, g, hh_cov, pop_cov, net_cov, lga_d, facilities, sync_lga_rows, sync_time_stats
    )

    conclusion = generate_narrative(
        _conclusion_prompt(cfg, g, hh_cov, pop_cov, net_cov, lga_d,
                            roster_ever_synced, roster_high),
        max_tokens=600,
    )
    _last = days_data[-1] if days_data else None
    slack_narrative = generate_narrative(
        _slack_prompt(cfg, g,
                      _last["cum_hh_visited"] if _last else g["hh_visited"],
                      _last["cum_pop_covered"] if _last else g["pop_covered"],
                      _last["cum_nets_distributed"] if _last else g["nets_distributed"],
                      campaign_days_for_target, elapsed_day,
                      roster_ever_synced, roster_high),
        max_tokens=400,
    )
    slack_text = (
        f"*{cfg['state_name']} — {cfg['campaign_name']}*\n\n"
        f"{slack_narrative}"
    )

    # Main report (full — includes the Data Quality section).
    doc = _build_doc(cfg, g=g, hh_cov=hh_cov, pop_cov=pop_cov, net_cov=net_cov,
                     lga_d=lga_d, facilities=facilities,
                     issues_data=issues_data, sync_lga_rows=sync_lga_rows,
                     sync_time_stats=sync_time_stats, sync_note=sync_note,
                     roster_ever_synced=roster_ever_synced, roster_high=roster_high,
                     days_data=days_data, chart_path=chart_path,
                     conclusion=conclusion, perf_link=perf_link, sync_link=sync_link,
                     perf_path=perf_path, sync_path=sync_path, partner=False,
                     dup_matrix=dup_matrix)
    out = cfg["docx_path"]
    doc.save(out)
    log.info(f"[report_itn] saved -> {out}")

    # Partner report — omits the Data Quality section and links a DQ-stripped Excel.
    # Same trigger condition as report.py: built whenever a partner channel is
    # configured, or always in cumulative mode.
    partner_docx_path = None
    if cfg.get("slack_channel_partners") or cfg.get("cumulative"):
        partner_perf_link = perf_link
        partner_perf_path = _make_partner_perf_xlsx_itn(perf_path)
        if partner_perf_path and not cfg.get("no_upload"):
            try:
                from dst_data_analysis_report.pipeline import notify as _notify
                p_title = f"{cfg['state_name']} ITN Performance Data (Partner) — {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                partner_perf_link = _notify.upload_file(
                    partner_perf_path, p_title, folder_id=_notify.campaign_folder_id(cfg))
                cfg["partner_perf_drive_link"] = partner_perf_link
            except Exception as e:
                log.warning(f"  partner perf Excel upload failed (non-fatal): {e}")
                partner_perf_link = ""

        partner_doc = _build_doc(cfg, g=g, hh_cov=hh_cov, pop_cov=pop_cov, net_cov=net_cov,
                                  lga_d=lga_d, facilities=facilities,
                                  issues_data=issues_data, sync_lga_rows=sync_lga_rows,
                                  sync_time_stats=sync_time_stats, sync_note=sync_note,
                                  roster_ever_synced=roster_ever_synced, roster_high=roster_high,
                                  days_data=days_data, chart_path=chart_path,
                                  conclusion=conclusion, perf_link=partner_perf_link, sync_link=sync_link,
                                  perf_path=partner_perf_path or perf_path, sync_path=sync_path, partner=True)
        partner_out = cfg["partner_docx_path"]
        partner_doc.save(partner_out)
        log.info(f"[report_itn] partner doc saved -> {partner_out}")
        partner_docx_path = partner_out

    return out, partner_docx_path, slack_text
