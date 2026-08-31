"""python-docx styling primitives shared by all Word report builders."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HDR_FILL  = "1A6496"
ALT_FILL  = "EBF3FB"
TITLE_RGB = RGBColor(0x17, 0x36, 0x5D)
GREY_RGB  = RGBColor(0x66, 0x66, 0x66)
FONT      = "Times New Roman"

STATUS_COLOR = {
    "HIGH":         RGBColor(0x1A, 0x7A, 0x1A),
    "MODERATE":     RGBColor(0xE0, 0x60, 0x00),
    "LOW":          RGBColor(0xCC, 0x00, 0x00),
    "NO TARGET":    RGBColor(0x88, 0x88, 0x88),
    "LOW ACTIVITY": RGBColor(0x88, 0x88, 0x88),
    "NOT REPORTED": RGBColor(0x88, 0x88, 0x88),
}

COV_FILL = {
    "HIGH":         "C6EFCE",
    "MODERATE":     "FFEB9C",
    "LOW":          "FFC7CE",
    "NO TARGET":    "F2F2F2",
    "LOW ACTIVITY": "F2F2F2",
    "NOT REPORTED": "F2F2F2",
}


def cov_band(pct):
    if pct >= 95:
        return "HIGH"
    if pct >= 70:
        return "MODERATE"
    return "LOW"


def add_hyperlink(para, text, url):
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)


def set_cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        borders.append(el)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def hdr(cell, text, size=9):
    set_cell_bg(cell, HDR_FILL)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def dat(cell, text, alt=False, bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=None):
    if alt:
        set_cell_bg(cell, ALT_FILL)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.alignment = align


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = False
    run.italic = True
    run.font.color.rgb = TITLE_RGB
    return p


def add_para(doc, text, style="Normal", size=None, color=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return p


def two_col_table(doc, rows_data, col_widths=(5, 9)):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for ri, (param, val) in enumerate(rows_data):
        hdr(table.cell(ri, 0), param, size=9)
        dat(table.cell(ri, 1), val, alt=(ri % 2 == 1), bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])
