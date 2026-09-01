"""
report.py — reads performance xlsx + CDD sync xlsx → Word .docx + Slack text
Claude writes the conclusion paragraph and the Slack summary.
"""
import logging
import os
from collections import defaultdict
from datetime import timedelta, datetime

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from dst_data_analysis_report.pipeline.core.llm import generate_narrative
from dst_data_analysis_report.pipeline.core.word import (
    ALT_FILL, FONT, GREY_RGB, STATUS_COLOR, TITLE_RGB,
    add_heading, add_hyperlink, add_para, cov_band, dat, hdr, set_cell_bg,
    set_cell_borders, two_col_table,
)

log = logging.getLogger(__name__)

# ── data loader ────────────────────────────────────────────────────────────────

def _load_performance_excel(path, drug_type):
    wb = openpyxl.load_workbook(path, read_only=True)
    if "ALL FACILITIES" not in wb.sheetnames:
        wb.close()
        raise FileNotFoundError(f"'ALL FACILITIES' tab missing in {path} — analyze.py may have failed mid-write")
    ws = wb["ALL FACILITIES"]
    rows_raw = [
        r for r in ws.iter_rows(min_row=3, values_only=True)
        if r[2] and str(r[2]).strip() not in ("", "TOTAL", "GRAND TOTAL")
    ]
    wb.close()

    lga_d = defaultdict(lambda: dict(
        hfs=0, target=0, treated=0, records=0,
        dups=0, missing_hh=0, missing_child=0, missing_gender=0,
        age59=0, age0=0,
        absent=0, refused=0, inelig=0, referred=0, died=0, migrated=0,
        drug1=0, drug2=0,
    ))
    facilities = []

    for row in rows_raw:
        # col order: #, lga, fac, tgt, rec, treated, not_treated, drug1, drug2, cov, status,
        #            absent, refused, inelig, referred, died, migrated, redose,
        #            age59, age0, missing_hh, missing_child, dups, del_com, wards
        num, lga, fac, tgt, rec, treated, not_treated, drug1, drug2, cov, status = row[:11]
        (absent, refused, inelig, referred, died, migrated, redose,
         age59, age0, missing_hh, missing_child, missing_gender, dups, del_com) = row[11:25]

        def i(v): return int(v or 0)

        if not lga:
            continue

        L = lga_d[str(lga).strip()]
        L["hfs"]          += 1
        L["target"]       += i(tgt)
        L["treated"]      += i(treated)
        L["records"]      += i(rec)
        L["dups"]         += i(dups)
        L["missing_hh"]      += i(missing_hh)
        L["missing_child"]   += i(missing_child)
        L["missing_gender"]  += i(missing_gender)
        L["age59"]        += i(age59)
        L["age0"]         += i(age0)
        L["absent"]       += i(absent)
        L["refused"]      += i(refused)
        L["inelig"]       += i(inelig)
        L["referred"]     += i(referred)
        L["died"]         += i(died)
        L["migrated"]     += i(migrated)
        L["drug1"]        += i(drug1)
        L["drug2"]        += i(drug2)

        facilities.append({
            "lga": str(lga).strip(), "fac": str(fac).strip(),
            "tgt": i(tgt), "rec": i(rec), "treated": i(treated),
            "cov": cov or "—", "status": str(status or "").strip(),
        })

    return lga_d, facilities


def _load_secondary_summary(path, spec):
    """
    Read one secondary-product tab (spec['tab']) from the performance Excel.
    Returns (total, {facility_name_lower: count}).
    """
    if not path or not spec or not os.path.exists(path):
        return 0, {}
    try:
        wb = openpyxl.load_workbook(path, read_only=True)
        if spec["tab"] not in wb.sheetnames:
            wb.close()
            return 0, {}
        ws    = wb[spec["tab"]]
        total = 0
        fac_d = {}
        for row in ws.iter_rows(min_row=3, values_only=True):
            if not row or not row[2]:
                continue
            fac = str(row[1] or "").strip()
            cnt = int(row[2] or 0)
            if fac and fac.upper() not in ("TOTAL", "GRAND TOTAL"):
                fac_d[fac.lower()] = cnt
                total += cnt
        wb.close()
        return total, fac_d
    except Exception as e:
        log.warning(f"[report] secondary summary load failed (non-fatal): {e}")
        return 0, {}


def _build_partner_performance_excel(perf_path):
    """
    Write a partner-safe copy of the performance Excel with data-quality columns
    stripped from every tab (Age>59, Age=0, Missing HH, Missing Child, Missing Gender,
    Duplicates) — matching the DQ sections the partner report already omits.
    Returns the new file path, or "" on failure.
    """
    from openpyxl.utils import get_column_letter
    DQ_COLS    = {"Age>59", "Age=0", "Missing HH", "Missing Child", "Missing Gender", "Duplicates"}
    HEADER_ROW = 2   # row 1 = banner, row 2 = column headers
    if not perf_path or not os.path.exists(perf_path):
        return ""
    try:
        wb = openpyxl.load_workbook(perf_path)
        for ws in wb.worksheets:
            headers = [c.value for c in ws[HEADER_ROW]]
            to_del  = sorted([i + 1 for i, h in enumerate(headers) if h in DQ_COLS], reverse=True)
            if not to_del:
                continue
            # unmerge the row-1 banner (spans all columns) before deleting, then re-merge
            merges = [str(m) for m in list(ws.merged_cells.ranges) if m.min_row == 1]
            for m in merges:
                ws.unmerge_cells(m)
            for ci in to_del:
                ws.delete_cols(ci, 1)
            if merges:
                ws.merge_cells(f"A1:{get_column_letter(ws.max_column)}1")
        base, ext = os.path.splitext(perf_path)
        out = f"{base}_partner{ext}"
        wb.save(out)
        log.info(f"[report] partner performance Excel written -> {out}")
        return out
    except Exception as e:
        log.warning(f"[report] partner perf Excel build failed (non-fatal): {e}")
        return ""


def _load_sync_summary(path):
    """Returns (lga_rows, time_stats) where time_stats = {label: (count, pct)}."""
    if not path or not os.path.exists(path):
        return [], {}
    wb   = openpyxl.load_workbook(path, read_only=True)
    ws   = wb["SUMMARY"]
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    wb.close()

    lga_rows   = []
    time_stats = {}
    for r in rows:
        if not r[0] and not r[1]:
            continue
        label = str(r[0] or "").strip()
        if "Synced by" in label:
            time_stats[label] = (r[1], r[2])   # (count, pct_str)
        elif r[1] and str(r[1]).strip().upper() not in ("GRAND TOTAL", "TOTAL"):
            lga_rows.append(r)
    return lga_rows, time_stats


def _load_facility_sync_rates(sync_path):
    """
    Read per-LGA tabs from the sync Excel and compute facility-level sync rates.
    Returns list of dicts sorted by sync_rate ascending (lowest first).
    """
    if not sync_path or not os.path.exists(sync_path):
        return []
    try:
        wb   = openpyxl.load_workbook(sync_path, read_only=True)
        skip = {"SUMMARY", "NEVER SYNCED", "LOW SYNCED"}
        fac_stats = {}   # facility -> {total, synced}

        for ws in wb.worksheets:
            if ws.title in skip:
                continue
            headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            try:
                fac_col    = headers.index("Health Facility")
                lga_col    = headers.index("LGA")
                status_col = headers.index("Status")
            except ValueError:
                continue

            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row or len(row) <= max(fac_col, lga_col, status_col):
                    continue
                fac    = str(row[fac_col] or "").strip()
                lga    = str(row[lga_col] or "").strip()
                status = str(row[status_col] or "").strip()
                if not fac:
                    continue
                key = (lga, fac)
                if key not in fac_stats:
                    fac_stats[key] = {"lga": lga, "fac": fac, "total": 0, "synced": 0}
                fac_stats[key]["total"] += 1
                if status != "NEVER SYNCED":
                    fac_stats[key]["synced"] += 1

        wb.close()
        results = []
        for v in fac_stats.values():
            rate = v["synced"] / v["total"] * 100 if v["total"] else 0
            results.append({**v, "rate": rate,
                             "rate_str": f"{rate:.1f}%",
                             "never": v["total"] - v["synced"]})
        return sorted(results, key=lambda x: x["rate"])   # lowest sync rate first
    except Exception as e:
        log.warning(f"[report] facility sync rate load failed (non-fatal): {e}")
        return []


# Days whose ES back-fill failed this run. Each one silently subtracts from the
# cumulative total the report leads with, so campaign_runner turns a non-empty
# list into a degraded outcome instead of publishing an understated number.
TRAJECTORY_FAILURES = []


def _load_daily_totals_from_es(cfg):
    """
    Cumulative trajectory straight from ES: per-day treated + records for EVERY
    campaign day (start .. cfg['DAY']), so the progress chart covers all days even
    when per-day performance Excel files are missing. Lightweight _count queries.
    """
    import requests, urllib3
    urllib3.disable_warnings()
    url, idx, auth = cfg["es_url"], cfg["ES_INDEX_TASK"], cfg["es_auth"]
    date_field = cfg.get("task_date_field", "taskDates")

    # campaign scope — mirror analyze._build_campaign_filters
    scope = []
    if cfg.get("task_campaign_filter"):
        if cfg.get("is_admin_console") and cfg.get("campaign_number"):
            scope.append({"term": {"Data.campaignNumber.keyword": cfg["campaign_number"]}})
        elif cfg.get("project_type_id"):
            scope.append({"term": {"Data.projectTypeId.keyword": cfg["project_type_id"]}})
        elif cfg.get("project_type"):
            scope.append({"term": {"Data.projectType.keyword": cfg["project_type"]}})
        if cfg.get("cycle_index"):
            scope.append({"term": {"Data.additionalDetails.cycleIndex.keyword": cfg["cycle_index"]}})

    min_age = 3 if cfg["drug_type"] == "SPAQ" else 1
    all_status = ["ADMINISTRATION_SUCCESS", "VISITED", "BENEFICIARY_INELIGIBLE",
                  "INELIGIBLE", "BENEFICIARY_REFERRED", "BENEFICIARY_DIED",
                  "BENEFICIARY_ABSENT", "BENEFICIARY_MIGRATED", "BENEFICIARY_REFUSED"]

    def _count(extra):
        q = {"query": {"bool": {"filter": scope + extra}}}
        r = requests.post(f"{url}/{idx}/_count", json=q, auth=auth, verify=False, timeout=60)
        r.raise_for_status()
        return r.json().get("count", 0)

    days = []
    for day_num in range(1, cfg["DAY"] + 1):
        d   = cfg["campaign_start"] + timedelta(days=day_num - 1)
        rng = {"range": {f"Data.{date_field}": {
            "gte": f"{d.isoformat()}T00:00:00.000Z", "lte": f"{d.isoformat()}T23:59:59.999Z"}}}
        try:
            # treated: mirror analyze._aggregate_batch is_treated exactly —
            # ADMINISTRATION_SUCCESS + age in [min_age,59] + quantity>=1 (+doseIndex if enabled).
            treated_f = [rng,
                         {"term":  {"Data.administrationStatus.keyword": "ADMINISTRATION_SUCCESS"}},
                         {"range": {"Data.age":      {"gte": min_age, "lte": 59}}},
                         {"range": {"Data.quantity": {"gte": 1}}}]
            if cfg.get("dose_index_filter"):
                treated_f.append({"term": {"Data.additionalDetails.doseIndex.keyword": "1"}})
            records = _count([rng, {"terms": {"Data.administrationStatus.keyword": all_status}}])
            treated = _count(treated_f)
        except Exception as e:
            # Recording 0 for a day we simply could not read UNDERSTATES the
            # cumulative figure that the Slack post leads with, and presents it
            # as if measured. Keep going (a partial trajectory beats none) but
            # register it so the run is marked degraded.
            TRAJECTORY_FAILURES.append(
                f"day {day_num} ({d.isoformat()}): {type(e).__name__}: {e}")
            log.error(f"[report] day {day_num} ES back-fill FAILED - that day "
                      f"counts as 0, so the cumulative total shown in the "
                      f"report and the Slack post is UNDERSTATED: {e}",
                      exc_info=True)
            records, treated = 0, 0
        days.append({"day": day_num, "date": d.strftime("%d %b"),
                     "records": records, "treated": treated,
                     "target": 0, "cov_pct": 0.0, "coverage": "N/A"})
    log.info(f"[report] cumulative trajectory from ES: {len(days)} days "
             f"(treated {sum(x['treated'] for x in days):,})")
    return days


def _load_daily_totals_from_files(cfg):
    """
    Read each day's performance Excel and return a list of daily totals dicts.
    Used for the day-by-day comparison table and bar chart.
    """
    days = []
    for day_num in range(1, cfg["DAY"] + 1):
        path = os.path.join(cfg["out_dir"], f"performance_day{day_num}.xlsx")
        if not os.path.exists(path):
            continue
        try:
            wb = openpyxl.load_workbook(path, read_only=True)
            ws = wb["ALL FACILITIES"]
            rows_raw = [
                r for r in ws.iter_rows(min_row=3, values_only=True)
                if r[2] and str(r[2]).strip() not in ("", "TOTAL", "GRAND TOTAL")
            ]
            wb.close()
            records = sum(int(r[4] or 0) for r in rows_raw)
            treated = sum(int(r[5] or 0) for r in rows_raw)
            target  = sum(int(r[3] or 0) for r in rows_raw)
            cov_pct = treated / target * 100 if target else 0
            date_label = (cfg["campaign_start"] + timedelta(days=day_num - 1)).strftime("%d %b")
            days.append({
                "day":      day_num,
                "date":     date_label,
                "records":  records,
                "treated":  treated,
                "target":   target,
                "cov_pct":  cov_pct,
                "coverage": f"{cov_pct:.1f}%" if target else "N/A",
            })
        except Exception as e:
            log.warning(f"Could not read performance_day{day_num}.xlsx: {e}")
    return days


def _generate_progress_chart(days_data, cfg, overall_target=None):
    """
    Cumulative coverage vs total campaign target bar chart.
    Formula: Cumulative Coverage % = Cumulative Treated (Days 1–N) ÷ Total Campaign Target × 100

    Daily mode:      Total Campaign Target = Daily Target × Campaign Days (per-day file targets).
    Cumulative mode: Total Campaign Target = overall_target (the full undivided campaign target),
                     so each day's bar shows cumulative treated climbing toward the overall target.
    """
    if not days_data:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        campaign_days = cfg["campaign_days"]
        labels        = [f"Day {d['day']}\n{d['date']}" for d in days_data]

        # Build cumulative coverage against total target.
        # Use a STABLE denominator for every bar — the fullest day's target sum
        # × campaign_days. Dividing cumulative treated by each day's OWN target sum
        # breaks on early/partial extracts: a day where few facilities have reported
        # yet has a tiny target sum, so cumulative-treated ÷ tiny-target explodes
        # (this is what produced the 916% Day-4 bar).
        cum_treated      = 0
        coverage         = []
        max_daily_target = max((d["target"] for d in days_data), default=0)
        for d in days_data:
            cum_treated += d["treated"]
            if cfg.get("cumulative") and overall_target:
                total_target = overall_target
            else:
                total_target = max_daily_target * campaign_days if max_daily_target else 0
            coverage.append(cum_treated / total_target * 100 if total_target else 0)

        def bar_color(c):
            if c >= 95: return "#1A7A1A"
            if c >= 70: return "#E06000"
            return "#CC0000"

        fig, ax = plt.subplots(figsize=(9, 4))
        fig.patch.set_facecolor("#F9F9F9")

        # Cap plotted bar height + label position to the axis so an outlier day
        # (e.g. a day whose target denominator is near-zero, producing an absurd %)
        # cannot shoot off the top and stretch the saved image into a tall white void.
        # Colour and the printed label still reflect the TRUE coverage value.
        Y_MAX     = 115
        BAR_CAP   = Y_MAX - 5     # 110 — keeps the value label inside the axis
        plot_vals = [min(c, BAR_CAP) for c in coverage]

        colors = [bar_color(c) for c in coverage]
        bars   = ax.bar(labels, plot_vals, color=colors, width=0.5, zorder=3)
        ax.set_title(f"{cfg['state_name']} — Cumulative Coverage vs Total Campaign Target",
                     fontsize=11, fontweight="bold", pad=10)
        ax.set_ylabel("Cumulative Coverage (%)", fontsize=9)
        ax.set_ylim(0, Y_MAX)
        ax.axhline(95, color="#1A7A1A", linestyle="--", linewidth=1, alpha=0.6, label="HIGH (95%)")
        ax.axhline(70, color="#E06000", linestyle="--", linewidth=1, alpha=0.6, label="MODERATE (70%)")
        ax.legend(fontsize=8)
        ax.grid(axis="y", linestyle="--", alpha=0.4, zorder=0)
        ax.set_axisbelow(True)
        for bar, val in zip(bars, coverage):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1.5,
                    f"{val:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold",
                    clip_on=True)

        plt.tight_layout(pad=2.0)
        chart_path = os.path.join(cfg["out_dir"], f"progress_chart_day{cfg['DAY']}.png")
        plt.savefig(chart_path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close()
        log.info(f"[report] progress chart saved -> {chart_path}")
        return chart_path
    except Exception as e:
        log.warning(f"[report] chart generation failed (non-fatal): {e}")
        return None


def _grand_totals(lga_d):
    g = defaultdict(int)
    for L in lga_d.values():
        for k, v in L.items():
            g[k] += v
    return g


# ── previous report reader ─────────────────────────────────────────────────────

def _read_previous_report(cfg):
    """
    Find the most recent earlier report .docx to use as context for Claude.

    Priority:
      1. Latest report from the same campaign day (earlier HHMM than now)
      2. If none, latest report from the previous campaign day (Day N-1)

    Returns the full paragraph text, or empty string if nothing found.
    """
    import glob as _glob

    state  = str(cfg.get("state_name", "")).strip().replace(" ", "_")
    day    = cfg["DAY"]
    folder = cfg["out_dir"]
    now_hm = datetime.now().strftime("%H%M")

    # 1. Same-day reports (earlier HHMM)
    pattern   = os.path.join(folder, f"{state}_Day{day}_Report_*.docx")
    same_day  = []
    for fp in sorted(_glob.glob(pattern)):
        tag = os.path.basename(fp).replace(f"{state}_Day{day}_Report_", "").replace(".docx", "")
        if tag.isdigit() and tag < now_hm:
            same_day.append(fp)

    if same_day:
        prev_path = same_day[-1]
    else:
        # 2. Fall back to last report of previous day
        prev_day     = day - 1
        prev_pattern = os.path.join(folder, f"{state}_Day{prev_day}_Report_*.docx")
        prev_day_files = sorted(_glob.glob(prev_pattern))
        if not prev_day_files:
            return ""
        prev_path = prev_day_files[-1]   # latest HHMM of previous day

    log.info(f"[report] reading previous report for context: {prev_path}")
    try:
        prev_doc = Document(prev_path)
        text = "\n".join(p.text for p in prev_doc.paragraphs if p.text.strip())
        return text
    except Exception as e:
        log.warning(f"[report] could not read previous report (non-fatal): {e}")
        return ""


def _extract_previous_conclusion(prev_report):
    """Extract only the conclusion paragraph from a previous report to minimise input tokens."""
    if not prev_report:
        return ""
    lines = [l.strip() for l in prev_report.split("\n") if l.strip()]
    for i, line in enumerate(lines):
        if line.startswith("6.") and "Conclusion" in line:
            snippet = " ".join(lines[i + 1: i + 5])
            return snippet[:400]
    return prev_report[:300]


def _conclusion_prompt(cfg, g, cov_pct, lga_d, sync_rows, sync_time_stats, prev_report):
    best_lga  = max(lga_d, key=lambda l: lga_d[l]["treated"] / lga_d[l]["target"] * 100
                    if lga_d[l]["target"] else 0, default="N/A")
    worst_lga = min(lga_d, key=lambda l: lga_d[l]["treated"] / lga_d[l]["target"] * 100
                    if lga_d[l]["target"] else 100, default="N/A")
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if len(r) > 2 and r[2])
    never       = sum(int(r[6] or 0) for r in sync_rows if len(r) > 6 and r[6])
    synced_cdds = (total_cdds - never) if cfg.get("cumulative") else \
                  sum(int(r[3] or 0) for r in sync_rows if len(r) > 3 and r[3])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    by17 = ""
    if sync_time_stats:
        for _, (count, pct) in sync_time_stats.items():
            by17 = f", {count:,} by 17:00 ({pct})"
    prev = _extract_previous_conclusion(prev_report)
    cum  = cfg.get("cumulative")
    header = (
        f"{cfg['campaign_name']} {cfg['state_name']} CUMULATIVE Days 1-{cfg['DAY']} "
        f"({cfg['START_LABEL']} to {cfg['END_LABEL']}) {cfg['drug_type']}"
        if cum else
        f"{cfg['campaign_name']} {cfg['state_name']} Day {cfg['DAY']}/{cfg['campaign_days']} {cfg['DATE_LABEL']} {cfg['drug_type']}"
    )

    return (
        f"{header}\n"
        f"Coverage:{cov_pct} Treated:{g['treated']:,} Target(overall):{g['target']:,} Records:{g['records']:,}\n"
        f"BestLGA:{best_lga} WorstLGA:{worst_lga}\n"
        f"Dups:{g['dups']:,} MissingChild:{g['missing_child']:,}\n"
        f"Sync:{synced_cdds:,}/{total_cdds:,}({sync_pct}){by17}\n"
        + (f"PrevReport:{prev}\n" if prev else "")
        + ("Write 5-sentence formal end-of-campaign conclusion covering the WHOLE campaign (all distribution + mop-up days): "
           if cum else "Write 5-sentence formal conclusion: ")
        + "(1)overall coverage vs full campaign target"
        + (" vs previous" if prev else "")
        + " (2)best LGA numbers (3)worst LGA implication (4)data quality action (5)sync outlook. Plain text only."
    )


def _issues_prompt(cfg, g, cov_pct, lga_d, facilities, sync_rows, sync_time_stats, prev_report):
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if len(r) > 2 and r[2])
    never       = sum(int(r[6] or 0) for r in sync_rows if len(r) > 6 and r[6])
    synced_cdds = (total_cdds - never) if cfg.get("cumulative") else \
                  sum(int(r[3] or 0) for r in sync_rows if len(r) > 3 and r[3])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    worst_sync  = sorted(
        [(str(r[1]), int(r[6] or 0), int(r[2] or 0))
         for r in sync_rows if r[1] and int(r[2] or 0) > 0],
        key=lambda x: -x[1],
    )
    worst_str = (f"{worst_sync[0][0]}:{worst_sync[0][1]}/{worst_sync[0][2]} never synced"
                 if worst_sync else "N/A")
    by17 = ""
    if sync_time_stats:
        for _, (count, pct) in sync_time_stats.items():
            by17 = f" by17:{count:,}({pct})"
    low_act = [f for f in facilities
               if f["rec"] < 10 and f["status"] != "NOT REPORTED"]
    prev = _extract_previous_conclusion(prev_report)
    is_last = cfg["DAY"] == cfg["campaign_days"]
    cum = cfg.get("cumulative")
    header = (
        f"{cfg['campaign_name']} {cfg['state_name']} CUMULATIVE Days 1-{cfg['DAY']} "
        f"({cfg['START_LABEL']} to {cfg['END_LABEL']}) FINAL"
        if cum else
        f"{cfg['campaign_name']} {cfg['state_name']} Day {cfg['DAY']}/{cfg['campaign_days']} {'FINAL' if is_last else ''} {cfg['DATE_LABEL']}"
    )

    return (
        f"{header}\n"
        f"Coverage:{cov_pct} Treated:{g['treated']:,}/{g['target']:,}\n"
        f"Sync:{synced_cdds:,}/{total_cdds:,}({sync_pct}) Never:{never:,}{by17} Worst:{worst_str}\n"
        f"LowActivity(<10):{len(low_act)} Dups:{g['dups']:,} MissingChild:{g['missing_child']:,}\n"
        + (f"PrevReport:{prev}\n" if prev else "")
        + 'Return ONLY a JSON array of 3-5 issues, no markdown:\n'
          '[{"observation":"...with numbers+LGA","status":"ACTIVE|RESOLVED","priority":"High|Moderate|Low",'
          '"notes":"actionable field instruction","data_type":"perf|sync"}]'
    )


def _generate_issues(cfg, g, cov_pct, lga_d, facilities, sync_rows, sync_time_stats, prev_report):
    import json
    prompt = _issues_prompt(cfg, g, cov_pct, lga_d, facilities, sync_rows, sync_time_stats, prev_report)
    raw = generate_narrative(prompt, max_tokens=1500)
    try:
        text = raw.strip()
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:])
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3]
        return json.loads(text)
    except Exception as e:
        # Salvage a truncated array: keep every complete {...} object up to the last one.
        try:
            t = text if text.startswith("[") else text[text.find("["):]
            cut = t.rfind("}")
            if cut != -1:
                salvaged = json.loads(t[:cut + 1].rstrip().rstrip(",") + "]")
                if isinstance(salvaged, list) and salvaged:
                    log.warning(f"[report] issues JSON truncated ({e}) — salvaged {len(salvaged)} complete issue(s)")
                    return salvaged
        except Exception:
            pass
        log.warning(f"[report] issues JSON parse failed: {e} — using raw text as single issue")
        return [{
            "observation": raw[:300] if raw else "No issues generated.",
            "status": "ACTIVE",
            "priority": "High",
            "notes": "Review manually.",
            "data_type": "perf",
        }]


def _slack_heading(cfg):
    """Deterministic one-line heading for the Slack post (state, campaign, day, date)."""
    if cfg.get("cumulative"):
        return (f"*{cfg['state_name']} — {cfg['campaign_name']} "
                f"(Cumulative Days 1-{cfg['DAY']}, {cfg['START_LABEL']} to {cfg['END_LABEL']})*")
    return (f"*{cfg['state_name']} — {cfg['campaign_name']} "
            f"(Day {cfg['DAY']} of {cfg['campaign_days']}, {cfg['DATE_LABEL']})*")


def _slack_prompt(cfg, g, cum_treated, docx_name, sync_rows, sync_time_stats, prev_report):
    """Prompt for a single flowing summary paragraph (figures woven into prose).

    Coverage in the Slack message is CUMULATIVE only:
    Cumulative Coverage % = Cumulative Treated (Days 1-N) / Total Campaign Target x 100.
    No daily target and no daily coverage are exposed to the narrative."""
    drug_type = cfg["drug_type"]
    d1 = "SPAQ2 (12-59 months)" if drug_type == "SPAQ" else "AZM 12-59 months"
    d2 = "SPAQ1 (3-11 months)"  if drug_type == "SPAQ" else "AZM 1-11 months"

    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if len(r) > 2 and r[2])
    never       = sum(int(r[6] or 0) for r in sync_rows if len(r) > 6 and r[6])
    synced_cdds = (total_cdds - never) if cfg.get("cumulative") else \
                  sum(int(r[3] or 0) for r in sync_rows if len(r) > 3 and r[3])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    by17 = ""
    if sync_time_stats:
        for _, (count, pct) in sync_time_stats.items():
            if count is not None:
                by17 = f"{count:,} ({pct})"
    not_admin = (g['absent'] + g['refused'] + g['inelig']
                 + g['referred'] + g['died'] + g['migrated'])
    prev = _extract_previous_conclusion(prev_report)
    cum  = cfg.get("cumulative")
    scope = (f"Cumulative Days 1-{cfg['DAY']} ({cfg['START_LABEL']} to {cfg['END_LABEL']})"
             if cum else
             f"Day {cfg['DAY']} of {cfg['campaign_days']} ({cfg['DATE_LABEL']})")

    # Cumulative coverage vs the TOTAL campaign target — the only coverage shown.
    total_target = g["target"] if cum else g["target"] * cfg["campaign_days"]
    treated_cum  = g["treated"] if cum else cum_treated
    cum_cov      = _coverage_str(treated_cum, total_target)

    facts = (
        f"Campaign: {cfg['campaign_name']} in {cfg['state_name']}, {scope}, {cfg['drug_type']}.\n"
        f"Cumulative coverage (Days 1-{cfg['DAY']} vs total campaign target): {cum_cov}\n"
        f"Cumulative children treated (Days 1-{cfg['DAY']}): {treated_cum:,}\n"
        f"Total campaign target: {total_target:,}\n"
        f"Children treated this day: {g['treated']:,}\n"
        f"Records submitted this day: {g['records']:,}\n"
        f"{d2}: {g['drug2']:,}\n"
        f"{d1}: {g['drug1']:,}\n"
        f"Not administered: {not_admin:,}\n"
        f"Duplicate records: {g['dups']:,}\n"
        f"Missing child name: {g['missing_child']:,}\n"
        f"CDDs synced: {synced_cdds:,} of {total_cdds:,} ({sync_pct})\n"
        f"Never synced: {never:,}\n"
        + (f"Synced by 17:00: {by17}\n" if by17 else "")
        + (f"Previous report summary: {prev}\n" if prev else "")
    )
    return (
        facts
        + "\nWrite a single flowing paragraph (4 to 6 sentences) for a Slack field update "
          "summarising the figures above. Weave the key numbers into the prose naturally. "
          "Coverage must be described ONLY as cumulative coverage against the total "
          "campaign target; never mention a daily target or daily coverage. "
          "Note the coverage trend"
        + (" versus the previous report" if prev else "")
        + ", then end with the single most urgent action for supervisors "
          "(prioritise data sync when sync coverage is low). "
          "Use the exact numbers given, do not invent any. Plain prose only — one paragraph, "
          "no heading, no bullet points, no emojis, no asterisks."
    )


# ── document builder ───────────────────────────────────────────────────────────

def _coverage_str(treated, target):
    if not target:
        return "N/A"
    return f"{treated/target*100:.1f}%"


def _lga_performance_table(doc, lga_d, lga_summary=True):
    header = ["LGA", "HFs", "Target", "Treated", "Coverage %", "Status"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (lga, L) in enumerate(sorted(lga_d.items()), 1):
        raw_pct = L["treated"] / L["target"] * 100 if L["target"] else None
        cov     = f"{raw_pct:.1f}%" if raw_pct is not None else "N/A"
        stat    = cov_band(raw_pct) if raw_pct is not None else "NO TARGET"
        vals    = [lga, L["hfs"], f"{L['target']:,}", f"{L['treated']:,}", cov, stat]
        row     = table.add_row()
        alt     = ri % 2 == 1
        for ci, val in enumerate(vals):
            if alt:
                set_cell_bg(row.cells[ci], ALT_FILL)
            set_cell_borders(row.cells[ci])
            p = row.cells[ci].paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.name = FONT; run.font.size = Pt(9)
            if ci in (4, 5):
                run.bold = True
                color = STATUS_COLOR.get(stat)
                if color: run.font.color.rgb = color
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER


def _dq_table(doc, lga_d):
    header = ["LGA", "Duplicates", "Missing HH", "Missing Child", "Age=0", "Age>59"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (lga, L) in enumerate(sorted(lga_d.items()), 1):
        vals = [lga, L["dups"], L["missing_hh"], L["missing_child"], L["age0"], L["age59"]]
        row  = table.add_row()
        alt  = ri % 2 == 1
        for ci, val in enumerate(vals):
            dat(row.cells[ci], val, alt=alt)


def _facility_performance_table(doc, facilities, perf_link="", secondary_specs=None, cumulative=False):
    """
    LOW coverage (<70%) and Low Activity (<10 records) facilities.
    Target Achievement % = Treated / Target * 100.
    Pop. Coverage        = Records / Target * 100.
    Treatment Coverage   = Treated / Records * 100.
    Target = overall campaign target in cumulative mode, daily target otherwise.
    Appends one count column per configured secondary product that has data.
    Cells colour-coded: >=95% green, 70-95% amber, <70% red.
    """
    tgt_word = "Campaign Target" if cumulative else "Daily Target"
    trt_word = "Treatment Coverage" if cumulative else "Daily Treatment Coverage"
    secondary_specs = secondary_specs or []
    show_specs = [s for s in secondary_specs
                  if any(f.get("secondary", {}).get(s["label"], 0) for f in facilities)]

    low_facs = [
        f for f in facilities
        if f["status"] in ("LOW", "LOW ACTIVITY") and f["tgt"] > 0
    ]
    low_facs = sorted(low_facs, key=lambda f: f["rec"] / f["tgt"] if f["tgt"] else 0)

    total_facs = len(facilities)
    shown      = min(len(low_facs), 20)
    note_text  = f"Showing {shown} LOW / Low Activity facilities of {total_facs} total. "
    note_p     = add_para(doc, note_text, size=9, color=GREY_RGB)
    if perf_link:
        add_hyperlink(note_p, "Full list in performance Excel ↗", perf_link)
    else:
        run = note_p.add_run("Full list in performance Excel.")
        run.font.name = FONT; run.font.size = Pt(9); run.font.color.rgb = GREY_RGB

    formula_lines = (
        f"Target Achievement % = Treated ÷ {tgt_word} × 100    |    "
        f"Pop. Coverage = Records ÷ {tgt_word} × 100    |    "
        f"{trt_word} = Treated ÷ Records × 100"
    )
    if show_specs:
        _parts = []
        for s in show_specs:
            _band = (f" age {s['age_min']}-{s['age_max']}m"
                     if s.get("age_min") is not None and s.get("age_max") is not None else "")
            _parts.append(f"{s['label']} = successful deliveries{_band}")
        formula_lines += "    |    " + ", ".join(_parts)
    add_para(doc, formula_lines, size=8, color=GREY_RGB)

    if not low_facs:
        add_para(doc, "No LOW or Low Activity facilities on this day.")
        return

    header = ["#", "District", "Health Facility", tgt_word, "Records",
              "Treated", "Not Treated", "Target Achievement %",
              "Pop. Coverage", trt_word]
    for s in show_specs:
        header.append(s["label"])

    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)

    for ri, f in enumerate(low_facs[:20], 1):
        achv_pct = f["treated"] / f["tgt"] * 100 if f["tgt"] else None
        pop_pct  = f["rec"]     / f["tgt"] * 100 if f["tgt"] else None
        trt_pct  = f["treated"] / f["rec"] * 100 if f["rec"] else None
        achv_str = f"{achv_pct:.1f}%" if achv_pct is not None else "N/A"
        pop_str  = f"{pop_pct:.1f}%"  if pop_pct  is not None else "N/A"
        trt_str  = f"{trt_pct:.1f}%"  if trt_pct  is not None else "N/A"
        not_trt  = f["rec"] - f["treated"]

        achv_band = cov_band(achv_pct) if achv_pct is not None else "NO TARGET"
        pop_band  = cov_band(pop_pct)  if pop_pct  is not None else "NO TARGET"
        trt_band  = cov_band(trt_pct)  if trt_pct  is not None else "NO TARGET"

        row = table.add_row()
        alt = ri % 2 == 1
        vals = [ri, f["lga"], f["fac"], f"{f['tgt']:,}", f"{f['rec']:,}",
                f"{f['treated']:,}", f"{not_trt:,}", achv_str, pop_str, trt_str]
        for s in show_specs:
            vals.append(f.get("secondary", {}).get(s["label"], 0))

        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            if alt:
                set_cell_bg(cell, ALT_FILL)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(9)
            if ci == 7:
                r.bold = True
                if STATUS_COLOR.get(achv_band): r.font.color.rgb = STATUS_COLOR[achv_band]
            if ci == 8:
                r.bold = True
                if STATUS_COLOR.get(pop_band): r.font.color.rgb = STATUS_COLOR[pop_band]
            if ci == 9:
                r.bold = True
                if STATUS_COLOR.get(trt_band): r.font.color.rgb = STATUS_COLOR[trt_band]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER


def _low_activity_table(doc, facilities):
    silent = [f for f in facilities if f["status"] == "NOT REPORTED"]
    if silent:
        add_para(doc, f"{len(silent):,} target-book facilities submitted no data "
                      f"at all this day (full list in the performance Excel).",
                 size=9, color=GREY_RGB)
    low_act = [f for f in facilities if f["rec"] < 10 and f["status"] != "NOT REPORTED"]
    if not low_act:
        add_para(doc, "No reporting facilities with fewer than 10 records on this day.")
        return
    header = ["Health Facility", "LGA", "Target", "Records", "Treated", "Status"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, f in enumerate(sorted(low_act, key=lambda x: x["rec"]), 1):
        row = table.add_row()
        alt = ri % 2 == 1
        dat(row.cells[0], f["fac"],     alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f["lga"],     alt=alt)
        dat(row.cells[2], f"{f['tgt']:,}", alt=alt)
        dat(row.cells[3], f"{f['rec']:,}", alt=alt)
        dat(row.cells[4], f"{f['treated']:,}", alt=alt)
        dat(row.cells[5], f["status"], alt=alt,
            color=STATUS_COLOR.get(f["status"]))


def _non_administration_table(doc, g):
    reasons = [
        ("Absent",      g["absent"]),
        ("Refused",     g["refused"]),
        ("Ineligible",  g["inelig"]),
        ("Referred",    g["referred"]),
        ("Died",        g["died"]),
        ("Migrated",    g["migrated"]),
    ]
    reasons = [(r, c) for r, c in reasons if c > 0]
    if not reasons:
        add_para(doc, "No non-administration events recorded.")
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr(table.cell(0, 0), "Reason")
    hdr(table.cell(0, 1), "Count")
    for ri, (reason, count) in enumerate(sorted(reasons, key=lambda x: -x[1]), 1):
        row = table.add_row()
        dat(row.cells[0], reason, alt=ri % 2 == 1, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f"{count:,}", alt=ri % 2 == 1)


def _dq_summary_table(doc, g):
    total = g["treated"] or 1
    metrics = [
        ("Duplicate Records",    g["dups"]),
        ("Missing HH Name",      g["missing_hh"]),
        ("Missing Child Name",   g["missing_child"]),
        ("Missing Gender",       g["missing_gender"]),
        ("Age = 0",              g["age0"]),
        ("Age > 59 months",      g["age59"]),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr(table.cell(0, 0), "Metric")
    hdr(table.cell(0, 1), "Count")
    hdr(table.cell(0, 2), "% of Treated")
    for ri, (metric, count) in enumerate(metrics, 1):
        row = table.add_row()
        pct = f"{count/total*100:.2f}%"
        dat(row.cells[0], metric, alt=ri % 2 == 1, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f"{count:,}", alt=ri % 2 == 1)
        dat(row.cells[2], pct, alt=ri % 2 == 1)

def _sync_table(doc, sync_rows, cfg, sync_time_stats=None):
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if len(r) > 2 and r[2])
    high        = sum(int(r[3] or 0) for r in sync_rows if len(r) > 3 and r[3])
    never       = sum(int(r[6] or 0) for r in sync_rows if len(r) > 6 and r[6])

    add_heading(doc, "5.1  FLW Sync Summary", 5)
    if cfg.get("cumulative"):
        # Whole-campaign totals. "Synced" = synced at least once over the campaign
        # (Total - Never = HIGH+MODERATE+LOW). HIGH alone means synced EVERY day, so
        # showing HIGH as "CDDs Synced" wrongly reads as 0 when no one synced all N days.
        synced_ever = total_cdds - never
        ever_pct    = f"{synced_ever/total_cdds*100:.1f}%" if total_cdds else "N/A"
        full_pct    = f"{high/total_cdds*100:.1f}%" if total_cdds else "N/A"
        summary_rows = [
            ("Total CDDs Registered",                        f"{total_cdds:,}"),
            (f"CDDs Synced (>=1 day, Days 1-{cfg['DAY']})",  f"{synced_ever:,}  ({ever_pct})"),
            (f"Fully Synced (all {cfg['DAY']} days)",        f"{high:,}  ({full_pct})"),
            ("Never Synced (whole campaign)",                f"{never:,}"),
            ("Report Period",  f"{cfg['START_LABEL']} to {cfg['END_LABEL']}"),
        ]
    else:
        sync_pct = f"{high/total_cdds*100:.1f}%" if total_cdds else "N/A"
        summary_rows = [
            ("Total CDDs Registered",   f"{total_cdds:,}"),
            ("CDDs Synced (total day)",  f"{high:,}  ({sync_pct})"),
            ("Never Synced",             f"{never:,}"),
        ]
        # Time-based breakdown (daily only — today-cutoff is meaningless cumulatively)
        if sync_time_stats:
            for label, (count, pct) in sync_time_stats.items():
                hour = label.replace("Synced by ", "").replace(" today (UTC)", "")
                if count is not None:
                    summary_rows.append(
                        (f"Synced by {hour}",
                         f"{count:,}  ({pct})  — early sync indicator")
                    )
        summary_rows += [("Report Date",  cfg["DATE_LABEL"])]
    two_col_table(doc, summary_rows, col_widths=(5, 6))
    doc.add_paragraph()

    if sync_rows:
        add_heading(doc, "5.1b  Sync Status by LGA", 5)
        cols = ["#", "LGA", "Total CDDs", "HIGH", "MODERATE", "LOW", "NEVER SYNCED", "% Never Synced"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for ci, h in enumerate(cols):
            hdr(table.cell(0, ci), h)
        for ri, row in enumerate(sync_rows, 1):
            tr = table.add_row()
            alt = ri % 2 == 1
            for ci in range(len(cols)):
                val = ri if ci == 0 else (row[ci - 1] if ci - 1 < len(row) else "")
                dat(tr.cells[ci], val, alt=alt)


# ── document builder ──────────────────────────────────────────────────────────

def _build_report_doc(cfg, *, g, cov_pct, lga_d, facilities, hfs_active, lgas_total,
               d1_label, d2_label, days_data, cum_records, cum_treated,
               cum_target, cum_cov, sync_rows, sync_time_stats, issues_data,
               conclusion, perf_link, sync_link, perf_path, sync_path,
               chart_path, secondary_specs=None, secondary_totals=None, partner=False):
    """Build and return a Document. partner=True omits DQ sections 3.2 and 3.5."""
    cum = cfg.get("cumulative", False)
    secondary_specs  = secondary_specs or []
    secondary_totals = secondary_totals or {}
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Title block — centred, navy, Times New Roman (matches tg_smc reference)
    title_p = doc.add_paragraph(style="Normal")
    title_p.clear()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"{cfg['state_name']}  —  {cfg['campaign_name']}")
    run.font.name = FONT; run.font.size = Pt(26)
    run.bold = True; run.font.color.rgb = TITLE_RGB

    sub = doc.add_paragraph(style="Normal")
    sub.clear()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_text = (
        f"{cfg['START_LABEL']} to {cfg['END_LABEL']}  —  Cumulative Campaign Report (Days 1 to {cfg['DAY']})"
        if cum else
        f"{cfg['START_LABEL']} to {cfg['END_LABEL']}  —  Day {cfg['DAY']} of {cfg['campaign_days']}"
    )
    r2 = sub.add_run(sub_text)
    r2.font.name = FONT; r2.font.size = Pt(14)
    r2.bold = True; r2.font.color.rgb = TITLE_RGB

    _period = (f"Cumulative Days 1-{cfg['DAY']}  —  {cfg['START_LABEL']} to {cfg['END_LABEL']}"
               if cum else f"Day {cfg['DAY']}  —  {cfg['DATE_LABEL']}")
    summary_line = (
        f"{_period}  |  "
        f"Extract: {cfg['DATE_LABEL']}, {datetime.now().strftime('%H:%M')}  |  "
        f"{hfs_active} LGAs  |  Coverage: {cov_pct}"
    )
    grey_p = add_para(doc, summary_line, size=8, color=GREY_RGB)
    grey_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Sync totals needed for overview table
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if len(r) > 2 and r[2])
    high_cdds   = sum(int(r[3] or 0) for r in sync_rows if len(r) > 3 and r[3])
    never_cdds  = sum(int(r[6] or 0) for r in sync_rows if len(r) > 6 and r[6])
    if cum:
        # Cumulative: "synced" = synced at least once over the whole campaign
        # (Total - Never = HIGH+MODERATE+LOW), NOT HIGH-only (which means every day).
        synced_cdds = total_cdds - never_cdds
    else:
        synced_cdds = high_cdds
    sync_pct_ov = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    by_17_val   = "—"
    if sync_time_stats:
        for label, (count, pct) in sync_time_stats.items():
            by_17_val = f"{count:,}  ({pct})" if count is not None else "—"
    not_admin = g['absent'] + g['refused'] + g['inelig'] + g['referred'] + g['died'] + g['migrated']

    # Section 1
    sec1_title = ("1.  Cumulative Campaign Overview" if cum
                  else f"1.  Day {cfg['DAY']} Operational Overview")
    add_heading(doc, sec1_title, 4)
    if cum:
        overview_rows = [
            ("State / Country",           cfg['state_name']),
            ("Activity",                  f"{cfg['drug_type']} Distribution  —  Cumulative (Days 1 to {cfg['DAY']})"),
            ("Campaign Dates",            f"{cfg['START_LABEL']} to {cfg['END_LABEL']}"),
            ("Data Extract Timestamp",    f"{cfg['DATE_LABEL']}, {datetime.now().strftime('%H:%M')}"),
            ("LGAs / Districts Covered",  f"{hfs_active} of {lgas_total}"),
            # ── Coverage vs overall campaign target ────────────────────────
            ("Overall Campaign Target",   f"{g['target']:,}"),
            ("Total Records Submitted",   f"{g['records']:,}"),
            ("Total Children Treated",    f"{g['treated']:,}"),
            ("Coverage vs Overall Target", cov_pct),
            # ── Drug split ─────────────────────────────────────────────────
            (d1_label,                    f"{g['drug1']:,}"),
            (d2_label,                    f"{g['drug2']:,}"),
            ("Not Administered",          f"{not_admin:,}"),
        ]
    else:
        # Cumulative block FIRST (headline numbers, measured against the TOTAL
        # campaign target — same basis as the Slack message), daily block second.
        total_campaign_target = g["target"] * cfg["campaign_days"]
        cum_cov_total = _coverage_str(cum_treated, total_campaign_target)
        overview_rows = [
            ("State / Country",               cfg['state_name']),
            ("Activity",                      f"{cfg['drug_type']} Distribution  —  Day {cfg['DAY']} of {cfg['campaign_days']}"),
            ("Date",                          cfg['DATE_LABEL']),
            ("Data Extract Timestamp",        f"{cfg['DATE_LABEL']}, {datetime.now().strftime('%H:%M')}"),
            ("Campaign Dates",                f"{cfg['START_LABEL']} to {cfg['END_LABEL']}"),
            ("LGAs / Districts Covered",      f"{hfs_active} of {lgas_total}"),
            # ── Cumulative vs total campaign target ────────────────────────────
            ("Total Campaign Target",         f"{total_campaign_target:,}"),
            (f"Cumulative Records (Days 1–{cfg['DAY']})",  f"{cum_records:,}"),
            (f"Cumulative Treated (Days 1–{cfg['DAY']})",  f"{cum_treated:,}"),
            (f"Cumulative Coverage (Days 1–{cfg['DAY']})", cum_cov_total),
            # ── This day ───────────────────────────────────────────────────────
            ("Daily Population Target",       f"{g['target']:,}"),
            ("Total Records Submitted",       f"{g['records']:,}"),
            ("Children Treated",              f"{g['treated']:,}"),
            ("Coverage vs Daily Target",      cov_pct),
            # ── Drug split ─────────────────────────────────────────────────────
            (d1_label,                        f"{g['drug1']:,}"),
            (d2_label,                        f"{g['drug2']:,}"),
            ("Not Administered",              f"{not_admin:,}"),
        ]
    for spec in secondary_specs:
        tot = secondary_totals.get(spec["label"], 0)
        if tot:
            band = (f" ({spec['age_min']}-{spec['age_max']}m)"
                    if spec.get("age_min") is not None and spec.get("age_max") is not None else "")
            overview_rows.append((f"{spec['label']} Distributed{band}", f"{tot:,}"))
    if cum:
        # ── Sync (whole-campaign totals; today-cutoff is meaningless here) ──
        overview_rows += [
            ("CDDs Registered",               f"{total_cdds:,}"),
            ("CDDs Synced (>=1 day)",         f"{synced_cdds:,}  ({sync_pct_ov})"),
            ("Never Synced",                  f"{never_cdds:,}"),
        ]
    else:
        overview_rows += [
            # ── Sync ───────────────────────────────────────────────────────────
            ("CDDs Registered",               f"{total_cdds:,}"),
            ("CDDs Synced",                   f"{synced_cdds:,}  ({sync_pct_ov})"),
            ("CDDs Synced by 17:00",          by_17_val),
        ]
    two_col_table(doc, overview_rows)
    doc.add_paragraph()

    # Coverage chart — directly under Section 1 overview table
    if chart_path and os.path.exists(chart_path):
        formula_p = add_para(
            doc,
            "Formula:  Cumulative Coverage % = Cumulative Treated (Days 1–N)  ÷  Total Campaign Target × 100"
            f"   |   Total Campaign Target = Daily Target × {cfg['campaign_days']} days",
            size=8, color=GREY_RGB,
        )
        formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(chart_path, width=Inches(5.5))
        doc.add_paragraph()

    # Section 2 — Issues log (Claude-generated, with Drive links per row)
    sec2_title = ("2.  Program Issues and Resolutions — Cumulative Log" if cum
                  else f"2.  Program Issues and Resolutions — Day {cfg['DAY']} Log")
    sec2_note  = (f"Whole campaign (Days 1 to {cfg['DAY']}). Issues identified across the extract and field reports."
                  if cum else
                  f"Day {cfg['DAY']} of {cfg['campaign_days']}. Issues identified during the extract and field reports.")
    add_heading(doc, sec2_title, 4)
    add_para(doc, sec2_note, size=9, color=GREY_RGB)

    issues_header = ["#", "Issue / Observation", "Status", "Priority", "Notes", "Data"]
    tbl2 = doc.add_table(rows=1, cols=len(issues_header))
    tbl2.style = "Table Grid"
    for ci, h in enumerate(issues_header):
        hdr(tbl2.cell(0, ci), h)

    _ACTIVE_RED   = RGBColor(0xCC, 0x00, 0x00)
    _RESOLVED_GRN = RGBColor(0x1A, 0x7A, 0x1A)
    _PRI_HIGH     = RGBColor(0xCC, 0x00, 0x00)
    _PRI_MOD      = RGBColor(0xE0, 0x60, 0x00)

    for ri, issue in enumerate(issues_data, 1):
        row2 = tbl2.add_row()
        alt  = ri % 2 == 1
        status   = str(issue.get("status",   "ACTIVE")).upper()
        priority = str(issue.get("priority", "High"))
        dtype    = str(issue.get("data_type","perf")).lower()

        dat(row2.cells[0], ri, alt=alt)
        dat(row2.cells[1], issue.get("observation", ""), alt=alt,
            align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        dat(row2.cells[2], status, alt=alt,
            color=_ACTIVE_RED if status == "ACTIVE" else _RESOLVED_GRN)
        dat(row2.cells[3], priority, alt=alt,
            color=_PRI_HIGH if priority == "High" else _PRI_MOD if priority == "Moderate" else None)
        dat(row2.cells[4], issue.get("notes", ""), alt=alt,
            align=WD_ALIGN_PARAGRAPH.LEFT, size=9)

        # Data column — hyperlink to relevant Drive file
        data_cell = row2.cells[5]
        if alt:
            set_cell_bg(data_cell, ALT_FILL)
        set_cell_borders(data_cell)
        dp = data_cell.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        link  = sync_link  if dtype == "sync"  else perf_link
        label = "CDD Sync Data" if dtype == "sync" else "Performance Data"
        fname = (os.path.basename(sync_path)  if dtype == "sync"  else os.path.basename(perf_path)) or label

        if link:
            add_hyperlink(dp, label + " ↗", link)
        else:
            run_d = dp.add_run(fname)
            run_d.font.name = FONT; run_d.font.size = Pt(8)
            run_d.font.color.rgb = GREY_RGB

    # column widths: #, Observation, Status, Priority, Notes, Data
    for ci, w in enumerate([Cm(0.8), Cm(5.5), Cm(1.5), Cm(1.6), Cm(4.5), Cm(2.5)]):
        for cell in tbl2.columns[ci].cells:
            cell.width = w
    doc.add_paragraph()

    # Section 3
    add_heading(doc, "3.  Distribution Data Analysis", 4)
    if not cum:
        add_para(doc,
                 f"All targets and coverage in this section are DAILY: "
                 f"Daily Target = Total Campaign Target ÷ Campaign Days "
                 f"({g['target'] * cfg['campaign_days']:,} ÷ {cfg['campaign_days']} "
                 f"= {g['target']:,} per day).",
                 size=9, color=GREY_RGB)

    add_heading(doc, "3.1  Performance by LGA", 5)
    _lga_performance_table(doc, lga_d)
    doc.add_paragraph()

    if not partner:
        add_heading(doc, "3.2  Data Quality by LGA", 5)
        _dq_table(doc, lga_d)
        doc.add_paragraph()

    fac_sec     = "3.3" if not partner else "3.2"
    nonadmin_sec = "3.4" if not partner else "3.3"

    add_heading(doc, f"{fac_sec}  Facility Performance Analysis", 5)
    add_para(doc, "LOW coverage (<70%) and Low Activity (<10 records) facilities. Sorted by Population Coverage ascending.", size=9, bold=True)
    _facility_performance_table(doc, facilities, perf_link=perf_link,
                    secondary_specs=secondary_specs, cumulative=cum)
    doc.add_paragraph()

    add_heading(doc, f"{nonadmin_sec}  Non-Administration Analysis", 5)
    _non_administration_table(doc, g)
    doc.add_paragraph()

    if not partner:
        add_heading(doc, "3.5  Data Quality Summary", 5)
        _dq_summary_table(doc, g)
        doc.add_paragraph()

    # Section 4 — Campaign Progress (day-by-day)
    add_heading(doc, f"4.  Campaign Progress  —  Days 1 to {cfg['DAY']}", 4)
    add_para(doc, "Coverage = Treated / Daily Target. Cumulative figures include all days to date.", size=9, color=GREY_RGB)

    if days_data:
        cols = ["Day", "Date", "Daily Target", "Records", "Treated", "Coverage"]
        tbl4 = doc.add_table(rows=1, cols=len(cols))
        tbl4.style = "Table Grid"
        for ci, h in enumerate(cols):
            hdr(tbl4.cell(0, ci), h)
        for ri, d in enumerate(days_data):
            row4 = tbl4.add_row()
            alt  = ri % 2 == 1
            cov_color = STATUS_COLOR.get(
                cov_band(d["cov_pct"]) if d["target"] else "NO TARGET"
            )
            vals = [f"Day {d['day']}", d["date"], f"{d['target']:,}",
                    f"{d['records']:,}", f"{d['treated']:,}", d["coverage"]]
            for ci, val in enumerate(vals):
                dat(row4.cells[ci], val, alt=alt,
                    color=cov_color if ci == 5 else None)
        # Cumulative totals row
        tot_row = tbl4.add_row()
        for ci, val in enumerate(["TOTAL", "", f"{cum_target:,}", f"{cum_records:,}", f"{cum_treated:,}", cum_cov]):
            dat(tot_row.cells[ci], val, bold=True)

    doc.add_paragraph()

    # Section 5 — Sync
    add_heading(doc, "5.  Health Facility Data Synchronisation Status", 4)
    _sync_table(doc, sync_rows, cfg, sync_time_stats)

    # 5.2 Top 10 facilities with lowest CDD sync rate
    fac_sync = _load_facility_sync_rates(sync_path)
    if fac_sync:
        doc.add_paragraph()
        add_heading(doc, "5.2  Top 10 Facilities with Lowest CDD Sync Rate", 5)
        note_p = add_para(doc, "Facilities requiring immediate follow-up from supervisors. ", size=9, color=GREY_RGB)
        if sync_link:
            add_hyperlink(note_p, "Full CDD sync data ↗", sync_link)
        low10  = fac_sync[:10]
        cols6  = ["#", "LGA", "Health Facility", "Total CDDs", "Synced", "Never Synced", "Sync Rate"]
        tbl6   = doc.add_table(rows=1, cols=len(cols6))
        tbl6.style = "Table Grid"
        for ci, h in enumerate(cols6):
            hdr(tbl6.cell(0, ci), h)
        for ri, f in enumerate(low10, 1):
            row6 = tbl6.add_row()
            alt  = ri % 2 == 1
            vals = [ri, f["lga"], f["fac"], f["total"], f["synced"], f["never"], f["rate_str"]]
            for ci, val in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if ci == 2 else WD_ALIGN_PARAGRAPH.CENTER
                dat(row6.cells[ci], val, alt=alt, align=align)

    # Section 6 — Conclusion (last section)
    doc.add_paragraph()
    add_heading(doc, "6.  Conclusion", 4)
    add_para(doc, conclusion, size=10)

    # Disclaimer — page footer (appears on every page)
    from datetime import timezone
    extracted_at = datetime.now(timezone.utc).strftime("%H:%M UTC")
    disclaimer = (
        f"Data extracted from DIGIT HCM as at {cfg['DATE_LABEL']}, {extracted_at}. "
        f"Figures are subject to change as field teams synchronise — refer to the DIGIT HCM dashboard for current data."
    )
    section = doc.sections[0]
    section.footer_distance = Cm(0.8)
    footer_para = section.footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(disclaimer)
    run.font.size = Pt(7)
    run.font.color.rgb = GREY_RGB
    run.font.name = FONT

    return doc


# ── public entry point ─────────────────────────────────────────────────────────

def _load_inputs(cfg):
    """Load both stage Excels plus secondary-product summaries into plain data."""
    perf_path = cfg["perf_xlsx"]
    sync_path = cfg["sync_xlsx"]
    if not os.path.exists(perf_path):
        raise FileNotFoundError(f"Performance Excel not found: {perf_path}")

    drug_type = cfg["drug_type"]
    lga_d, facilities          = _load_performance_excel(perf_path, drug_type)
    # NOT REPORTED zero rows (LGA "—") must stay in the TOTALS — that is the
    # standard-target denominator — but must NOT reach display/narrative
    # surfaces: the pseudo-group is permanently 0% and would e.g. always win
    # "worst LGA" in the conclusion. Tables and prompts use this filtered copy.
    lga_display = {k: v for k, v in lga_d.items() if k != "—"}
    sync_rows, sync_time_stats = _load_sync_summary(sync_path)

    secondary_specs  = cfg.get("secondary_products", []) or []
    secondary_totals = {}
    for spec in secondary_specs:
        tot, fac_d = _load_secondary_summary(perf_path, spec)
        secondary_totals[spec["label"]] = tot
        if fac_d:
            for f in facilities:
                f.setdefault("secondary", {})[spec["label"]] = fac_d.get(f["fac"].lower(), 0)

    g = _grand_totals(lga_d)
    inputs = {
        "perf_path":        perf_path,
        "sync_path":        sync_path,
        "lga_d":            lga_d,
        "lga_display":      lga_display,
        "facilities":       facilities,
        "sync_rows":        sync_rows,
        "sync_time_stats":  sync_time_stats,
        "secondary_specs":  secondary_specs,
        "secondary_totals": secondary_totals,
        "g":                g,
        "cov_pct":          _coverage_str(g["treated"], g["target"]),
        "hfs_active":       len({f["lga"] for f in facilities
                                 if f["status"] != "NOT REPORTED"}),
        "lgas_total":       cfg.get("lgas_total") or len(lga_display),
        "d1_label":         "SPAQ2 (12-59m)" if drug_type == "SPAQ" else "AZM 12-59m",
        "d2_label":         "SPAQ1 (3-11m)"  if drug_type == "SPAQ" else "AZM 1-11m",
    }
    log.info(f"[report:load] {len(facilities)} facilities, {len(lga_d)} LGAs, "
             f"coverage {inputs['cov_pct']}")
    return inputs


def _build_trajectory(cfg, g):
    """Day-by-day totals and the progress chart for the cumulative view."""
    if cfg.get("cumulative"):
        # Core numbers come from the cumulative Excel (g); the trajectory is built
        # day-by-day from ES so it covers EVERY campaign day regardless of which
        # per-day performance files exist on disk.
        days_data = _load_daily_totals_from_es(cfg)
        if not days_data:
            days_data = _load_daily_totals_from_files(cfg)
        daily_tgt = round(g["target"] / cfg["campaign_days"]) if cfg.get("campaign_days") else 0
        for d in days_data:
            d["target"]   = daily_tgt
            d["cov_pct"]  = d["treated"] / daily_tgt * 100 if daily_tgt else 0.0
            d["coverage"] = f"{d['cov_pct']:.1f}%" if daily_tgt else "N/A"
        return {
            "days_data":   days_data,
            "cum_records": g["records"],
            "cum_treated": g["treated"],
            "cum_target":  g["target"],
            "cum_cov":     _coverage_str(g["treated"], g["target"]),
            "chart_path":  _generate_progress_chart(days_data, cfg, overall_target=g["target"]),
        }

    # ES first, exactly as the cumulative branch does. out_dir is a disposable
    # per-run temp dir under Airflow, so days 1..N-1 never exist on disk: reading
    # only files made cum_treated equal to TODAY, and that value is what the
    # Slack post presents as "Cumulative children treated (Days 1-N)".
    days_data   = _load_daily_totals_from_es(cfg) or _load_daily_totals_from_files(cfg)
    cum_treated = sum(d["treated"] for d in days_data)
    cum_target  = sum(d["target"]  for d in days_data)
    return {
        "days_data":   days_data,
        "cum_records": sum(d["records"] for d in days_data),
        "cum_treated": cum_treated,
        "cum_target":  cum_target,
        "cum_cov":     f"{cum_treated/cum_target*100:.1f}%" if cum_target else "N/A",
        "chart_path":  _generate_progress_chart(days_data, cfg),
    }


def _report_period(cfg):
    return (f"Cumulative Days 1-{cfg['DAY']}" if cfg.get("cumulative")
            else f"Day {cfg['DAY']}")


def _publish_excels(cfg, perf_path, sync_path):
    """Upload the stage Excels to Drive so report tables can hyperlink them.

    Skipped when no_upload is set (links then show plain filenames). Links are
    stashed on cfg so callers such as the cumulative runner can list them.
    """
    perf_link = ""
    sync_link = ""
    if cfg.get("no_upload"):
        log.info("[report:publish] no_upload set — skipping Drive upload of Excels")
    else:
        try:
            from dst_data_analysis_report.pipeline import notify
            fid    = notify.campaign_folder_id(cfg)
            now_hm = datetime.now().strftime("%H:%M")
            period = _report_period(cfg)
            log.info("[report:publish] uploading performance Excel to Drive ...")
            perf_link = notify.upload_file(
                perf_path,
                f"{cfg['state_name']} {period} Performance Data — {cfg['DATE_LABEL']} {now_hm}",
                folder_id=fid)
            if sync_path and os.path.exists(sync_path):
                log.info("[report:publish] uploading CDD sync Excel to Drive ...")
                sync_link = notify.upload_file(
                    sync_path,
                    f"{cfg['state_name']} {period} CDD Sync Data — {cfg['DATE_LABEL']} {now_hm}",
                    folder_id=fid)
        except Exception as e:
            log.warning(f"[report:publish] Drive upload of Excels failed (non-fatal): {e}")

    cfg["perf_drive_link"] = perf_link
    cfg["sync_drive_link"] = sync_link
    return perf_link, sync_link


def _publish_partner_excel(cfg, partner_perf_path):
    """Upload the DQ-stripped partner Excel. Returns its Drive link ("" on skip/failure)."""
    if cfg.get("no_upload"):
        return ""
    try:
        from dst_data_analysis_report.pipeline import notify
        title = (f"{cfg['state_name']} {_report_period(cfg)} Performance Data (Partner) — "
                 f"{cfg['DATE_LABEL']} {datetime.now().strftime('%H:%M')}")
        link = notify.upload_file(partner_perf_path, title,
                                  folder_id=notify.campaign_folder_id(cfg))
        cfg["partner_perf_drive_link"] = link
        return link
    except Exception as e:
        log.warning(f"[report:publish] partner perf Excel upload failed (non-fatal): {e}")
        return ""


def _generate_narratives(cfg, inputs, trajectory, prev_report):
    """All LLM calls in one place: issues log, conclusion paragraph, Slack text."""
    g, cov_pct = inputs["g"], inputs["cov_pct"]
    sync_rows, sync_time_stats = inputs["sync_rows"], inputs["sync_time_stats"]

    log.info("[report:narrative] generating issues log ...")
    issues_data = _generate_issues(cfg, g, cov_pct, inputs["lga_display"], inputs["facilities"],
                                   sync_rows, sync_time_stats, prev_report)
    log.info(f"[report:narrative] {len(issues_data)} issues generated")

    log.info("[report:narrative] generating conclusion ...")
    conclusion = generate_narrative(
        _conclusion_prompt(cfg, g, cov_pct, inputs["lga_display"], sync_rows, sync_time_stats, prev_report),
        max_tokens=600)

    log.info("[report:narrative] generating Slack text ...")
    slack_narrative = generate_narrative(
        _slack_prompt(cfg, g, trajectory["cum_treated"], os.path.basename(cfg["docx_path"]),
                      sync_rows, sync_time_stats, prev_report),
        max_tokens=400)
    slack_text = _slack_heading(cfg) + "\n\n" + slack_narrative

    return issues_data, conclusion, slack_text


def _render_docs(cfg, render_params):
    """Build and save the internal doc, and the partner doc when configured."""
    doc = _build_report_doc(cfg, **render_params, partner=False)
    out = cfg["docx_path"]
    doc.save(out)
    log.info(f"[report:render] saved -> {out}")

    # Partner report (omits DQ sections and links a DQ-stripped Excel). Built when
    # a partner channel is configured, and always in cumulative mode.
    partner_docx_path = None
    if cfg.get("slack_channel_partners") or cfg.get("cumulative"):
        partner_render = dict(render_params)
        partner_perf_path = _build_partner_performance_excel(render_params["perf_path"])
        if partner_perf_path:
            partner_render["perf_path"] = partner_perf_path
            partner_render["perf_link"] = _publish_partner_excel(cfg, partner_perf_path)
        partner_doc = _build_report_doc(cfg, **partner_render, partner=True)
        partner_docx_path = cfg["partner_docx_path"]
        partner_doc.save(partner_docx_path)
        log.info(f"[report:render] partner doc saved -> {partner_docx_path}")

    return out, partner_docx_path


def run(cfg):
    log.info(f"[report] {cfg['state_name']} Day {cfg['DAY']} ...")

    inputs     = _load_inputs(cfg)
    trajectory = _build_trajectory(cfg, inputs["g"])

    prev_report = _read_previous_report(cfg)
    if prev_report:
        log.info(f"  previous report loaded ({len(prev_report)} chars)")
    else:
        log.info("  no previous report — first extract today")

    perf_link, sync_link = _publish_excels(cfg, inputs["perf_path"], inputs["sync_path"])
    issues_data, conclusion, slack_text = _generate_narratives(cfg, inputs, trajectory, prev_report)

    render_params = dict(
        g=inputs["g"], cov_pct=inputs["cov_pct"], lga_d=inputs["lga_display"],
        facilities=inputs["facilities"],
        hfs_active=inputs["hfs_active"], lgas_total=inputs["lgas_total"],
        d1_label=inputs["d1_label"], d2_label=inputs["d2_label"],
        days_data=trajectory["days_data"], cum_records=trajectory["cum_records"],
        cum_treated=trajectory["cum_treated"], cum_target=trajectory["cum_target"],
        cum_cov=trajectory["cum_cov"],
        sync_rows=inputs["sync_rows"], sync_time_stats=inputs["sync_time_stats"],
        issues_data=issues_data, conclusion=conclusion,
        perf_link=perf_link, sync_link=sync_link,
        perf_path=inputs["perf_path"], sync_path=inputs["sync_path"],
        chart_path=trajectory["chart_path"],
        secondary_specs=inputs["secondary_specs"],
        secondary_totals=inputs["secondary_totals"],
    )

    out, partner_docx_path = _render_docs(cfg, render_params)
    return out, partner_docx_path, slack_text
